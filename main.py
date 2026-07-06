"""
Punto Rojo — API de procesamiento de facturas XML DIAN
Servidor FastAPI que recibe ZIP/XML y devuelve los datos procesados.
"""
import os, io, re, zipfile, json, difflib
import xml.etree.ElementTree as ET
from datetime import datetime, date
from decimal import Decimal, ROUND_HALF_UP
from typing import Optional
from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from supabase import create_client, Client
from dotenv import load_dotenv
from collections import Counter

load_dotenv()

app = FastAPI(title="Punto Rojo API", version="1.0.0")

FRONTEND_ORIGIN = os.getenv("FRONTEND_ORIGIN", "https://costos-punto-rojo-app.vercel.app")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[FRONTEND_ORIGIN, "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

SUPABASE_URL = os.getenv("SUPABASE_URL", "")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY", "")
IVA_DEFAULT  = 19.0
MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB

# ─── Endpoints públicos que no requieren sesión ───────────────────────────────
PUBLIC_PATHS = {"/", "/catalogo", "/categorias", "/subcategorias", "/todas-subcategorias"}

# ─── Endpoints destructivos/sensibles que requieren rol admin ─────────────────
# Estos endpoints solo deben poder ejecutarse si el usuario autenticado tiene
# rol "admin" en user_profiles. El frontend ya oculta los botones a usuarios
# sin permiso, pero esto evita que alguien llame la API directamente (ej. con
# Postman o las herramientas de desarrollador) usando un token válido de un
# rol distinto a admin.
ADMIN_ONLY_PATHS = {
    "/limpiar-tablas",
    "/config-admin",
    "/unificar-productos",
}

# Endpoints que requieren admin O contabilidad (igual que el resto de la app,
# donde contabilidad tiene permisos casi iguales a admin salvo Admin general)
ADMIN_OR_CONTABILIDAD_PREFIXES = (
    "/eliminar-factura/",
)

def get_supabase() -> Client:
    return create_client(SUPABASE_URL, SUPABASE_KEY)

def get_user_from_token(token: str):
    """Devuelve el objeto user de Supabase si el token es una sesión activa, o None."""
    if not token:
        return None
    try:
        sb = get_supabase()
        result = sb.auth.get_user(token)
        if result is not None and result.user is not None:
            return result.user
    except Exception:
        pass
    return None

def verify_session_token(token: str) -> bool:
    """Verifica que el token sea una sesión activa en Supabase."""
    return get_user_from_token(token) is not None

def get_user_rol(email: str) -> str:
    """Consulta user_profiles y devuelve el rol del usuario, o '' si no existe."""
    if not email:
        return ""
    try:
        sb = get_supabase()
        res = sb.table("user_profiles").select("rol").eq("email", email).maybe_single().execute()
        return (res.data.get("rol") or "") if res.data else ""
    except Exception:
        return ""

@app.middleware("http")
async def auth_middleware(request: Request, call_next):
    path = request.url.path
    # Permitir OPTIONS (preflight CORS) y rutas públicas sin token
    if request.method == "OPTIONS" or path in PUBLIC_PATHS:
        return await call_next(request)
    # Verificar token en header Authorization: Bearer <token>
    auth_header = request.headers.get("Authorization", "")
    token = auth_header.replace("Bearer ", "").strip() if auth_header.startswith("Bearer ") else ""
    user = get_user_from_token(token)
    if user is None:
        return JSONResponse(status_code=401, content={"detail": "No autorizado"})

    requiere_solo_admin = path in ADMIN_ONLY_PATHS
    requiere_admin_o_contabilidad = any(path.startswith(p) for p in ADMIN_OR_CONTABILIDAD_PREFIXES)

    if requiere_solo_admin or requiere_admin_o_contabilidad:
        email = (user.email or "").lower()
        rol = get_user_rol(email)
        roles_permitidos = {"admin"} if requiere_solo_admin else {"admin", "contabilidad"}
        if rol not in roles_permitidos:
            return JSONResponse(status_code=403, content={"detail": "Esta acción requiere permisos de administrador"})

    return await call_next(request)

# ─── Helpers XML ────────────────────────────────────────────────────────────

def local_name(tag: str) -> str:
    return tag.split("}")[-1] if "}" in tag else tag

def all_descendants(parent, name: str):
    for el in parent.iter():
        if local_name(el.tag) == name:
            yield el

def first_text(parent, path_names: list) -> str:
    if not path_names:
        return parent.text.strip() if parent.text else ""
    for child in parent:
        if local_name(child.tag) == path_names[0]:
            return first_text(child, path_names[1:])
    return ""

def parse_decimal(text, default=0.0) -> float:
    try:
        return float(str(text or "").replace(",", "").strip())
    except Exception:
        return float(default)

def money(value) -> float:
    try:
        d = Decimal(str(float(value or 0)))
        return float(d.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))
    except Exception:
        return 0.0

def extract_invoice_xml(raw_xml: bytes) -> ET.Element:
    text = raw_xml.decode("utf-8", errors="replace")
    cdata = re.search(r'<!\[CDATA\[(.*?)\]\]>', text, re.DOTALL)
    inner = cdata.group(1) if cdata else text
    try:
        return ET.fromstring(inner.encode("utf-8"))
    except ET.ParseError:
        return ET.fromstring(re.sub(r'<\?xml[^>]*\?>', '', inner).strip().encode("utf-8"))

# ─── Detección de presentación ───────────────────────────────────────────────

def detect_packaging(desc: str):
    if not desc:
        return 1, 1, 1
    d = desc.upper()

    m = re.search(r'CJ\s*X?\s*(\d+)', d)
    if m:
        cj_num = int(m.group(1))
        pm = re.search(r'(?:PC|PQ|PK)\s*X?\s*(\d+)', d)
        if pm:
            up = int(pm.group(1))
            pc = max(1, cj_num // up)
            return up, cj_num, pc
        pos_cj = m.start()
        texto_antes = d[:pos_cj]
        matches_x = re.findall(r'(?<!\d)X\s*(\d+)', texto_antes)
        if matches_x:
            up = int(matches_x[-1])
            pc = cj_num
            uc = up * pc
            return up, uc, pc
        return 1, cj_num, cj_num

    m = re.search(r'(?:PC|PQ|PK)\s*X?\s*(\d+)', d)
    if m:
        up = int(m.group(1))
        return up, up, 1

    m = re.search(r'\bC(\d+)\b', d)
    if m:
        packs = int(m.group(1))
        pm = re.search(r'(?<!\d)X\s*(\d+)', d)
        up = int(pm.group(1)) if pm else 1
        return up, up * packs, packs

    return 1, 1, 1

def default_sale_flags(pres: str, up: int, packs: int, box: int):
    """Devuelve los 7 flags de venta (unidad, paquete, caja, millar, kg, rollo, metro)
    según la presentación detectada/asignada."""
    vu = vp = vc = vmillar = vkg = vrollo = vmetro = False
    if pres == "Caja/Paca" or pres == "Caja":
        vu = up > 1
        vp = packs > 1
        vc = True
    elif pres == "Paquete":
        vu = up > 1
        vp = True
    elif pres == "Millar":
        vmillar = True
    elif pres == "Kg":
        vkg = True
    elif pres == "Rollo":
        vrollo = True
    elif pres == "Metro":
        vmetro = True
    else:
        vu = True
    return vu, vp, vc, vmillar, vkg, vrollo, vmetro

def get_pres_sugerida(up: int, box: int, packs: int) -> str:
    if box > 1 and packs > 1:
        return "Caja/Paca"
    if up > 1:
        return "Paquete"
    return "Unidad"

# ─── Parser principal ────────────────────────────────────────────────────────

def parse_invoice(root: ET.Element) -> dict:
    supplier = (
        first_text(root, ["AccountingSupplierParty","Party","PartyName","Name"]) or
        first_text(root, ["AccountingSupplierParty","Party","PartyLegalEntity","RegistrationName"]) or
        "PROVEEDOR SIN NOMBRE"
    )
    proveedor_nit = (
        first_text(root, ["AccountingSupplierParty","Party","PartyTaxScheme","CompanyID"]) or
        first_text(root, ["AccountingSupplierParty","Party","PartyLegalEntity","CompanyID"])
    )
    number   = first_text(root, ["ID"]) or f"SIN_NUMERO_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    fecha    = first_text(root, ["IssueDate"]) or str(date.today())
    cufe     = ""
    for u in all_descendants(root, "UUID"):
        if u.text and u.text.strip():
            cufe = u.text.strip(); break

    subtotal = parse_decimal(first_text(root, ["LegalMonetaryTotal","LineExtensionAmount"]))
    total    = parse_decimal(first_text(root, ["LegalMonetaryTotal","PayableAmount"]))
    iva      = money(total - subtotal) if subtotal and total and total >= subtotal else 0.0

    lines = []
    def _id_valido(txt):
        txt = (txt or "").strip()
        return txt if txt and txt != "0" else ""

    # Primero recolectar el SKU y el nombre de CADA línea, sin decidir nada todavía
    raw_skus_por_linea = []
    descs_por_linea = []
    qtys_por_linea = []
    precios_por_linea = []
    for line in all_descendants(root, "InvoiceLine"):
        raw_sku = (
            _id_valido(first_text(line, ["Item","StandardItemIdentification","ID"])) or
            _id_valido(first_text(line, ["Item","SellersItemIdentification","ID"])) or
            ""
        )
        raw_skus_por_linea.append(raw_sku)
        descs_por_linea.append(first_text(line, ["Item","Description"]) or first_text(line, ["Item","Name"]) or "")
        qtys_por_linea.append(parse_decimal(first_text(line, ["InvoicedQuantity"]), 1) or 1)
        precios_por_linea.append(parse_decimal(first_text(line, ["Price","PriceAmount"])))

    # SKUs que aparecen en más de una línea son inútiles (proveedor los usa como placeholder,
    # ej. el mismo código base para dos lotes/despachos distintos del mismo producto)
    sku_counts = Counter(s for s in raw_skus_por_linea if s)
    skus_duplicados = {s for s, c in sku_counts.items() if c > 1}

    # Nombres de líneas que van a terminar con SKU temporal -- ya sea porque no traían
    # ningún código, O porque el código que traían resultó ser un placeholder repetido
    # entre varias líneas. Si ese nombre se repite entre varias de esas líneas, son
    # productos distintos que chocarían en la unique constraint
    # (proveedor_id, sku_proveedor, nombre_factura) si se guardan todos con sku_proveedor="".
    nombres_sin_sku = [
        desc for raw_sku, desc in zip(raw_skus_por_linea, descs_por_linea)
        if not raw_sku or raw_sku in skus_duplicados
    ]
    nombre_counts = Counter(nombres_sin_sku)
    nombres_duplicados_sin_sku = {n for n, c in nombre_counts.items() if c > 1 and n}

    # De esos nombres duplicados sin SKU, cuáles tienen ADEMÁS la misma cantidad y precio
    # en todas sus líneas -- eso ya no es el caso legítimo de PRECORTE (varios rollos,
    # mismo nombre, distinto peso/cantidad cada uno), sino muy probablemente la MISMA
    # compra contada dos veces por error. Se usa para avisar, no para fusionar solo.
    grupos_qty_precio: dict = {}
    for nombre, qty, precio in zip(descs_por_linea, qtys_por_linea, precios_por_linea):
        if nombre in nombres_duplicados_sin_sku:
            grupos_qty_precio.setdefault(nombre, set()).add((round(qty, 3), round(precio, 2)))
    nombres_posible_linea_repetida = {n for n, valores in grupos_qty_precio.items() if len(valores) == 1}

    for i, line in enumerate(all_descendants(root, "InvoiceLine"), start=1):
        qty      = parse_decimal(first_text(line, ["InvoicedQuantity"]), 1) or 1
        desc     = first_text(line, ["Item","Description"]) or first_text(line, ["Item","Name"]) or ""

        def _id_valido(txt):
            txt = (txt or "").strip()
            return txt if txt and txt != "0" else ""

        raw_sku = (
            _id_valido(first_text(line, ["Item","StandardItemIdentification","ID"])) or
            _id_valido(first_text(line, ["Item","SellersItemIdentification","ID"])) or
            ""
        )
        # Si el SKU está duplicado en varias líneas, es un placeholder inútil → temporal
        if raw_sku and raw_sku not in skus_duplicados:
            sku = raw_sku
        else:
            sku = f"L{i:03d}"
        price    = parse_decimal(first_text(line, ["Price","PriceAmount"]))
        line_ext = parse_decimal(first_text(line, ["LineExtensionAmount"]))

        # Descuento de línea (AllowanceCharge con ChargeIndicator=false).
        # BaseAmount es el valor ANTES del descuento; Amount es el descuento;
        # por estándar DIAN/UBL siempre se cumple: BaseAmount - Amount = LineExtensionAmount.
        desc_amt  = 0.0
        desc_pct  = 0.0
        base_amt  = 0.0
        for ac in all_descendants(line, "AllowanceCharge"):
            charge_ind = first_text(ac, ["ChargeIndicator"])
            if charge_ind.lower() == "false":
                desc_amt = parse_decimal(first_text(ac, ["Amount"]))
                desc_pct = parse_decimal(first_text(ac, ["MultiplierFactorNumeric"]))
                base_amt = parse_decimal(first_text(ac, ["BaseAmount"]))

        # precio_unitario_factura: precio BRUTO por unidad, antes de cualquier descuento de línea.
        # - Si hay descuento de línea, usamos BaseAmount/qty (consistente y confiable entre
        #   proveedores; PriceAmount no lo es: algunos lo reportan por caja, otros con otro
        #   significado que no cuadra matemáticamente con BaseAmount/LineExtensionAmount).
        # - Si NO hay descuento, mantenemos el cálculo original: LineExtensionAmount/qty
        #   (precio pre-IVA, consistente con subtotal/total e iva_detectado). En algunos
        #   proveedores (ej: Sidell) PriceAmount ya incluye IVA aunque la factura global sea
        #   NO_INCLUIDO, lo que duplicaría el IVA si lo usáramos directamente.
        if desc_amt > 0 and base_amt > 0:
            precio_unitario = base_amt / qty if qty else base_amt
        else:
            precio_unitario = (line_ext / qty) if qty else price
            if precio_unitario <= 0:
                precio_unitario = price

        line_iva = 0.0
        iva_pct  = IVA_DEFAULT
        for ta in all_descendants(line, "TaxTotal"):
            line_iva += parse_decimal(first_text(ta, ["TaxAmount"]))
            pct = parse_decimal(first_text(ta, ["TaxSubtotal","TaxCategory","Percent"]), 0)
            if pct: iva_pct = pct

        up, box, packs = detect_packaging(desc)
        pres = get_pres_sugerida(up, box, packs)
        vu, vp, vc, vmillar, vkg, vrollo, vmetro = default_sale_flags(pres, up, packs, box)

        lines.append({
            "linea": i,
            "sku_proveedor": str(sku or ""),
            "nombre_factura": desc,
            "nombre_duplicado_sin_sku": ((not raw_sku) or (raw_sku in skus_duplicados)) and (desc in nombres_duplicados_sin_sku),
            "posible_linea_repetida": desc in nombres_posible_linea_repetida,
            "cantidad_facturada": qty,
            "precio_unitario_factura": precio_unitario,
            "subtotal_linea": line_ext,
            "iva_linea": line_iva,
            "iva_porcentaje": iva_pct,
            "nombre_punto_rojo": desc,
            "categoria": "",
            "presentacion_facturada": pres,
            "unidades_por_paquete": up,
            "paquetes_por_caja": packs,
            "unidades_por_caja": up * packs,
            "venta_unidad": vu,
            "venta_paquete": vp,
            "venta_caja": vc,
            "venta_millar": vmillar,
            "venta_kg": vkg,
            "venta_rollo": vrollo,
            "venta_metro": vmetro,
            "markup_unidad_pct": 40.0,
            "markup_paquete_pct": 35.0,
            "markup_caja_pct": 31.58,
            "markup_millar_pct": 40.0,
            "markup_kg_pct": 40.0,
            "markup_rollo_pct": 40.0,
            "markup_metro_pct": 40.0,
            "descuento_factura_pct": desc_pct,
            "descuento_factura_amt": desc_amt,
            "nota_descuento": f"Descuento en factura {desc_pct:.0f}% (${(desc_amt / max(qty, 1)):,.0f} por unidad compra)" if desc_amt > 0 else "",
        })

    if not subtotal: subtotal = sum(x["subtotal_linea"] for x in lines)
    if not iva:      iva = sum(x["iva_linea"] for x in lines)
    if not total:    total = subtotal + iva

    retefuente_xml = 0.0
    for ta in all_descendants(root, "TaxTotal"):
        tax_id   = ""
        tax_name = ""
        for sc in all_descendants(ta, "TaxScheme"):
            tax_id   = first_text(sc, ["ID"]) or ""
            tax_name = (first_text(sc, ["Name"]) or "").upper()
        if tax_id in ("06", "05", "07") or "RETE" in tax_name or "RTEFUENTE" in tax_name:
            retefuente_xml += parse_decimal(first_text(ta, ["TaxAmount"]))

    inpusu = 0.0
    for note_el in all_descendants(root, "Note"):
        note_text = (note_el.text or "").upper()
        if "INPUSU" in note_text or "PLASTICO" in note_text:
            m = re.search(r'\$\s*([\d,\.]+)', note_text)
            if m:
                try:
                    # Formato "$ 53,432.00" -> coma = separador de miles, punto = decimal
                    inpusu = float(m.group(1).replace(",", ""))
                except Exception:
                    pass

    tolerancia = max(500, total * 0.02)
    diff_sin_inpusu = abs((subtotal + iva) - total)
    diff_con_inpusu = abs((subtotal + iva) - (total - inpusu))
    # Solo restar el INPUSU si efectivamente acerca el total (evita falsos "INCLUIDO"
    # cuando el INPUSU es solo informativo y no está sumado al total)
    if inpusu > 0 and diff_con_inpusu < diff_sin_inpusu:
        total_ajustado = total - inpusu
    else:
        total_ajustado = total

    iva_detectado = (
        "NO_INCLUIDO"
        if abs((subtotal + iva) - total_ajustado) < tolerancia
        else "INCLUIDO"
    )

    # Forma de pago desde XML
    payment_code = first_text(root, ["PaymentMeans", "PaymentMeansCode"]) or ""
    due_date = first_text(root, ["PaymentMeans", "PaymentDueDate"]) or ""
    issue_date = first_text(root, ["IssueDate"]) or ""
    CONTADO_CODES = {"10", "48", "49"}
    if payment_code.strip() in CONTADO_CODES:
        forma_pago_xml = "CONTADO"
    elif due_date and issue_date and due_date > issue_date:
        forma_pago_xml = "CREDITO"
    elif payment_code.strip() == "1" and (not due_date or due_date == issue_date):
        forma_pago_xml = "CONTADO"
    else:
        forma_pago_xml = "CREDITO" if payment_code.strip() else ""

    return {
        "proveedor": supplier,
        "proveedor_nit": proveedor_nit,
        "cufe": cufe,
        "numero_factura": number,
        "fecha": fecha,
        "subtotal_factura": money(subtotal),
        "iva_factura": money(iva),
        "total_factura": money(total),
        "retefuente_xml": money(retefuente_xml),
        "inpusu": money(inpusu),
        "iva_detectado": iva_detectado,
        "forma_pago_xml": forma_pago_xml,
        "lineas": lines,
    }

# ─── Cálculo de costos y precios ─────────────────────────────────────────────

def cost_without_tax(precio: float, iva_mode: str, iva_pct: float = IVA_DEFAULT) -> float:
    if iva_mode == "INCLUIDO":
        return money(precio / (1 + iva_pct / 100))
    return money(precio)

def calc_costs(costo_fact: float, pres: str, up: int, pc: int, precio_es_por: str = "", unidades_por_millar: int = 1000):
    up = max(up, 1); pc = max(pc, 1)
    uc = up * pc
    upm = max(int(unidades_por_millar or 1000), 1)

    base = precio_es_por if precio_es_por else (
        "Caja"    if pres == "Caja/Paca" else
        "Paquete" if pres == "Paquete"   else
        "Unidad"
    )

    if base == "Caja":
        cc = costo_fact
        cp = money(cc / pc) if pc > 1 else cc
        cu = money(cp / up) if up > 1 else cp
    elif base == "Millar":
        # El precio facturado es por millar (ej: 10 paquetes de 100, o 20 de 50 -> upm = 1000)
        cc = costo_fact
        cu = money(cc / upm) if upm > 1 else cc
        cp = money(cu * up) if up > 1 else cu
    elif base == "Paquete":
        cp = costo_fact
        cu = money(cp / up) if up > 1 else cp
        cc = money(cp * pc)
    else:
        # Unidad, Kg, Rollo, Metro -> el precio facturado es directamente el costo de 1 unidad/kg/rollo/metro
        cu = costo_fact
        cp = money(cu * up)
        cc = money(cu * uc)

    return cu, cp, cc

def inferir_precio_es_por(precio: float, up: int, pc: int, uc: int) -> str:
    """Infiere si el precio facturado es por unidad, paquete o caja
    usando lógica matemática: el costo unitario resultante debe ser razonable (>= 0.5 COP).
    
    Regla principal: si precio / up >= 0.5 y up > 1 → precio es por paquete.
    Si precio / uc >= 0.5 pero precio / up < 0.5 → precio es por caja.
    """
    MIN_CU = 0.5
    up = max(up, 1); uc = max(uc, 1)

    cu_paquete = precio / up if up > 1 else None
    cu_caja    = precio / uc if uc > up else None

    # Si precio/up da costo razonable y up > 1 → precio es por paquete
    if cu_paquete and cu_paquete >= MIN_CU:
        if cu_caja is None or cu_caja < MIN_CU:
            return "Paquete"
        if up > 1:
            return "Paquete"

    # Si precio/uc razonable y precio/up muy alto o no razonable → precio es por caja
    if cu_caja and cu_caja >= MIN_CU and (cu_paquete is None or cu_paquete > 10000):
        return "Caja"

    return "Unidad"

def sale_price(costo: float, markup: float, transporte: float = 0.0) -> float:
    costo_total = (costo + transporte) * (1 + IVA_DEFAULT / 100)
    m = float(markup or 0)
    if m >= 100:
        return money(costo_total * (m / 100))
    m_frac = max(0.0, m / 100.0)
    return money(costo_total / (1 - m_frac))

def add_calcs(lines: list, iva_mode: str) -> list:
    result = []
    for l in lines:
        up   = max(int(l.get("unidades_por_paquete") or 1), 1)
        pc   = max(int(l.get("paquetes_por_caja") or 1), 1)
        uc   = max(up * pc, 1)
        pres = l.get("presentacion_facturada", "Unidad")
        mu   = float(l.get("markup_unidad_pct") or 40)
        mp   = float(l.get("markup_paquete_pct") or 35)
        mc   = float(l.get("markup_caja_pct") or 31.58)

        precio_fact = l.get("precio_unitario_factura") or (l.get("subtotal_linea", 0) / max(l.get("cantidad_facturada", 1), 1))
        desc_pct = float(l.get("descuento_factura_pct") or 0)
        if l.get("descuento_afecta_costo") and desc_pct > 0:
            precio_fact = money(precio_fact * (1 - desc_pct / 100))

        costo_base       = cost_without_tax(precio_fact, iva_mode, l.get("iva_porcentaje", IVA_DEFAULT))
        transporte       = float(l.get("transporte_adicional", 0) or 0)
        costo_fact_final = costo_base + transporte

        # Usar precio_es_por del usuario si viene, si no inferir automáticamente
        precio_es_por = l.get("precio_es_por") or ""
        if not precio_es_por:
            precio_es_por = inferir_precio_es_por(costo_fact_final, up, pc, uc)

        upm = int(l.get("unidades_por_millar") or 1000)
        cu, cp, cc = calc_costs(costo_fact_final, pres, up, pc, precio_es_por, upm)

        row = {**l,
            "precio_es_por": precio_es_por,
            "unidades_por_millar": upm,
            "unidades_por_caja": uc,
            "costo_unidad_sin_iva": cu,
            "costo_paquete_sin_iva": cp,
            "costo_caja_sin_iva": cc,
            "precio_unidad_con_iva": sale_price(cu, mu),
            "precio_paquete_con_iva": sale_price(cp, mp),
            "precio_caja_con_iva": sale_price(cc, mc),
        }
        result.append(row)
    return result

# ─── Búsqueda de productos similares en Supabase ─────────────────────────────

def _es_sku_temporal(sku: str) -> bool:
    """SKUs tipo L001, L002... son temporales asignados cuando el proveedor
    no manda un código real. Nunca deben usarse para buscar matches."""
    return bool(re.match(r'^L\d{3}$', (sku or "").strip()))

def nombres_muy_distintos(nombre_nuevo: str, nombre_guardado: str) -> bool:
    """Compara la descripcion de la linea nueva contra la del producto ya guardado
    cuando el emparejamiento fue por codigo de proveedor (no por nombre). Sirve de
    aviso -- no bloquea nada -- para detectar proveedores que reutilizan el mismo
    codigo para variantes distintas (ej: 'SIN CORDON' vs 'CON CORDON')."""
    a = (nombre_nuevo or "").strip().upper()
    b = (nombre_guardado or "").strip().upper()
    if not a or not b:
        return False
    # Regla especial: una dice "CON X" y la otra "SIN X" -- casi siempre son
    # variantes distintas del mismo producto base, aunque el resto del texto
    # sea casi identico (por eso la similitud general no las atrapa).
    if (" CON " in f" {a} " and " SIN " in f" {b} ") or (" SIN " in f" {a} " and " CON " in f" {b} "):
        return True
    # Similitud general de texto -- descripciones del mismo producto varian poco
    # entre facturas (mayusculas, espacios); una caida grande sugiere otro producto.
    ratio = difflib.SequenceMatcher(None, a, b).ratio()
    return ratio < 0.6

def find_similar_product(supabase: Client, proveedor_nit: str, sku: str, nombre: str, nombre_duplicado_sin_sku: bool = False):
    SELECT_MATCH_COLS = (
        "id,sku_interno,sku_proveedor,nombre_factura,nombre_punto_rojo,categoria,subcategoria,"
        "presentacion_facturada,precio_es_por,unidades_por_paquete,paquetes_por_caja,unidades_por_caja,"
        "costo_unidad_sin_iva,markup_unidad_pct,markup_paquete_pct,markup_caja_pct,"
        "markup_millar_pct,markup_kg_pct,markup_rollo_pct,markup_metro_pct,"
        "multiplicador_rollo,multiplicador_metro,"
        "venta_unidad,venta_paquete,venta_caja,venta_millar,venta_kg,venta_rollo,venta_metro,costo_transporte,es_materia_prima"
    )
    # Resolver el proveedor_id una sola vez -- se necesita en ambos caminos de abajo
    # para no emparejar nunca productos de un proveedor con el codigo de otro.
    prov_id = None
    if proveedor_nit:
        try:
            prov = supabase.table("proveedores").select("id").eq("nit", proveedor_nit).limit(1).execute()
            if prov.data:
                prov_id = prov.data[0]["id"]
        except Exception:
            pass

    # SKU vacío o temporal → nunca hacer match por SKU
    if not sku or _es_sku_temporal(sku):
        # Si el nombre se repite en otra línea de ESTA MISMA factura sin SKU real, no
        # emparejar automaticamente: son productos distintos y deben quedar separados
        # (ej. un codigo generico reutilizado para variantes distintas). PERO si ademas
        # ya existe un producto activo de este proveedor con el mismo nombre (de una
        # factura anterior o de otra linea), avisar -- puede ser el mismo producto
        # comprado en varios rollos/lotes dentro de esta factura (ver caso IPM
        # PRECORTE 31*50: 3 lineas, mismo codigo y nombre, 3 rollos con distinto peso
        # en kg -- sin este aviso se crean 3 productos duplicados en vez de 1).
        if nombre_duplicado_sin_sku or prov_id is None:
            posible_duplicado = None
            if prov_id is not None and nombre_duplicado_sin_sku:
                try:
                    rdup = supabase.table("productos").select(
                        "id,sku_interno,nombre_punto_rojo,costo_unidad_sin_iva"
                    ).eq("proveedor_id", prov_id).eq("nombre_factura", nombre).eq("activo", True).limit(1).execute()
                    if rdup.data:
                        posible_duplicado = rdup.data[0]
                except Exception:
                    pass
            return {"match": "Nuevo", "producto": None, "posible_duplicado": posible_duplicado}
        # Compra recurrente de un producto sin código de proveedor: buscar si ya existe
        # uno de este proveedor con el mismo nombre y sku_proveedor="" para actualizarlo
        # en vez de crear un duplicado que choque contra la unique constraint.
        try:
            r = supabase.table("productos").select(SELECT_MATCH_COLS) \
                .eq("proveedor_id", prov_id).eq("sku_proveedor", "") \
                .eq("nombre_factura", nombre).eq("activo", True).limit(1).execute()
            if r.data:
                return {"match": "Exacto", "producto": r.data[0], "via": "nombre"}
        except Exception:
            pass
        return {"match": "Nuevo", "producto": None}

    # Match exacto por SKU del proveedor -- SIEMPRE filtrado por proveedor_id, para que
    # un codigo corto/generico (ej. "1", "2") de un proveedor nunca empareje con el
    # producto de OTRO proveedor que use el mismo codigo por coincidencia.
    if prov_id is None:
        return {"match": "Nuevo", "producto": None}
    try:
        r = supabase.table("productos").select(SELECT_MATCH_COLS) \
            .eq("sku_proveedor", sku).eq("proveedor_id", prov_id).eq("activo", True).limit(1).execute()
        if r.data:
            return {"match": "Exacto", "producto": r.data[0], "via": "sku"}
    except Exception:
        pass
    # El codigo no emparejo con ningun producto de este proveedor -- puede ser que el
    # proveedor haya cambiado su esquema de codigos entre facturas (ej. MIO paso de
    # codigos cortos a codigos de barras y genero duplicados en el catalogo). Avisar si
    # ya existe un producto activo de este proveedor con el mismo nombre de factura.
    posible_duplicado = None
    try:
        rdup = supabase.table("productos").select(
            "id,sku_interno,nombre_punto_rojo,costo_unidad_sin_iva"
        ).eq("proveedor_id", prov_id).eq("nombre_factura", nombre).eq("activo", True).limit(1).execute()
        if rdup.data:
            posible_duplicado = rdup.data[0]
    except Exception:
        pass
    return {"match": "Nuevo", "producto": None, "posible_duplicado": posible_duplicado}

def check_duplicate(supabase: Client, cufe: str, numero_factura: str) -> bool:
    try:
        if cufe:
            r = supabase.table("facturas").select("id").eq("cufe", cufe).execute()
            if r.data:
                return True
        if numero_factura:
            r2 = supabase.table("facturas").select("id").eq("numero_factura", numero_factura).execute()
            return bool(r2.data)
        return False
    except Exception:
        return False

def get_proveedor_info(supabase: Client, nit: str, nombre: str) -> dict:
    try:
        if nit:
            r = supabase.table("proveedores_contables").select(
                "id,proveedor_nombre,forma_pago,descuento_pct,aplica_retefuente,tipo,regimen,descuento_afecta_costo"
            ).eq("nit", nit).limit(1).execute()
            if r.data:
                return {**r.data[0], "encontrado_por": "nit"}

        r2 = supabase.rpc("match_proveedor", {"nombre_buscar": nombre}).execute()
        if r2.data:
            row = r2.data[0]
            return {
                "id":                    row.get("id"),
                "proveedor_nombre":      row.get("nombre") or row.get("proveedor_nombre"),
                "nit":                   row.get("nit"),
                "forma_pago":            row.get("forma_pago"),
                "descuento_pct":         row.get("descuento_pct"),
                "aplica_retefuente":     row.get("aplica_retefuente"),
                "tipo":                  row.get("tipo"),
                "regimen":               row.get("regimen"),
                "descuento_afecta_costo": row.get("descuento_afecta_costo", False),
                "similitud":             row.get("similitud"),
                "encontrado_por":        "nombre",
            }
    except Exception:
        pass
    return {}

def estimar_retefuente(supabase: Client, prov_info: dict, subtotal: float, fecha_factura: str, retefuente_xml: float = 0.0) -> float:
    """Calcula el retefuente esperado para una factura, igual lógica que /save-invoice.
    Se usa tanto en el preview (/parse-invoice) como al guardar, para evitar que ambos
    lugares calculen valores distintos."""
    if retefuente_xml and retefuente_xml > 0:
        return retefuente_xml
    aplica_rete = prov_info.get("aplica_retefuente", "NO")
    if aplica_rete != "SI":
        return 0.0
    try:
        fecha_factura = fecha_factura or str(date.today())
        # vigente_desde / vigente_hasta pueden quedar vacíos cuando el parámetro está
        # activo sin fecha de inicio/corte definida — un NULL en cualquiera de las dos
        # columnas no debe excluir el registro.
        params = supabase.table("parametros_retefuente").select(
            "porcentaje,base_minima"
        ).eq("aplica_a", "COMPRAS").eq("activo", True) \
         .or_(f"vigente_desde.is.null,vigente_desde.lte.{fecha_factura}") \
         .or_(f"vigente_hasta.is.null,vigente_hasta.gte.{fecha_factura}").limit(1).execute()
        if params.data:
            pct_rete = float(params.data[0].get("porcentaje") or 2.5)
            base_min = float(params.data[0].get("base_minima") or 1148000)
            if subtotal >= base_min:
                return round(subtotal * pct_rete / 100, 2)
    except Exception:
        pass
    return 0.0

# ─── Endpoints ───────────────────────────────────────────────────────────────

@app.get("/")
def root():
    return {"status": "ok", "service": "Punto Rojo API", "version": "1.0.0"}

@app.post("/parse-invoice")
async def parse_invoice_endpoint(
    file: UploadFile = File(...),
    iva_mode_override: Optional[str] = None,
):
    raw = await file.read()
    if len(raw) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="Archivo demasiado grande. Máximo 10 MB.")
    filename = file.filename or ""

    pdf_base64 = ""
    try:
        if filename.lower().endswith(".zip"):
            with zipfile.ZipFile(io.BytesIO(raw)) as z:
                xml_names = [n for n in z.namelist() if n.lower().endswith(".xml") and not n.lower().startswith("__")]
                if not xml_names:
                    raise HTTPException(status_code=400, detail="El ZIP no contiene archivos XML.")
                raw_xml = z.read(xml_names[0])
                pdf_names = [n for n in z.namelist() if n.lower().endswith(".pdf")]
                if pdf_names:
                    import base64 as b64lib
                    pdf_base64 = b64lib.b64encode(z.read(pdf_names[0])).decode("utf-8")
        elif filename.lower().endswith(".xml"):
            raw_xml = raw
        else:
            raise HTTPException(status_code=400, detail="Solo se aceptan archivos XML o ZIP.")

        root_el = extract_invoice_xml(raw_xml)
        invoice  = parse_invoice(root_el)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Error al leer el XML: {str(e)}")

    iva_mode = iva_mode_override or invoice["iva_detectado"]

    try:
        sb = get_supabase()
        invoice["es_duplicado"] = check_duplicate(sb, invoice["cufe"], invoice["numero_factura"])
        prov_info = get_proveedor_info(sb, invoice["proveedor_nit"] or "", invoice["proveedor"])
        invoice["proveedor_info"] = prov_info
        invoice["retefuente_estimado"] = estimar_retefuente(
            sb, prov_info,
            float(invoice.get("subtotal_factura") or 0),
            invoice.get("fecha") or "",
            float(invoice.get("retefuente_xml") or 0),
        )
        for line in invoice["lineas"]:
            line["descuento_afecta_costo"] = False
    except Exception as e:
        invoice["supabase_error"] = str(e)
        invoice["es_duplicado"] = False
        invoice["proveedor_info"] = {}
        invoice["retefuente_estimado"] = float(invoice.get("retefuente_xml") or 0)

    invoice["iva_mode_usado"] = iva_mode

    try:
        sb = get_supabase()
        for line in invoice["lineas"]:
            match_info = find_similar_product(sb, invoice["proveedor_nit"] or "", line["sku_proveedor"], line["nombre_factura"], line.get("nombre_duplicado_sin_sku", False))
            line["match_tipo"]  = match_info["match"]
            line["producto_bd"] = match_info["producto"]

            # Aviso no bloqueante: mismo nombre/proveedor que un producto ya existente,
            # detectado aunque el match haya quedado en "Nuevo" (codigo duplicado
            # dentro de esta factura). No cambia el guardado -- solo avisa para que
            # se revise si hay que unificar en vez de crear un producto repetido.
            posible_dup = match_info.get("posible_duplicado")
            if posible_dup:
                line["posible_duplicado_existente"] = posible_dup

            if match_info["match"] == "Exacto" and match_info["producto"]:
                p = match_info["producto"]
                line["producto_id"]            = p.get("id")
                line["nombre_punto_rojo"]      = p.get("nombre_punto_rojo") or line["nombre_factura"]
                line["categoria"]              = p.get("categoria") or ""
                line["subcategoria"]           = p.get("subcategoria") or ""
                line["presentacion_facturada"] = p.get("presentacion_facturada") or line["presentacion_facturada"]
                line["precio_es_por"]          = p.get("precio_es_por") or ""
                line["unidades_por_paquete"]   = p.get("unidades_por_paquete") or line["unidades_por_paquete"]
                line["paquetes_por_caja"]      = p.get("paquetes_por_caja") or line["paquetes_por_caja"]
                up_match = p.get("unidades_por_paquete") or line["unidades_por_paquete"]
                pc_match = p.get("paquetes_por_caja") or line["paquetes_por_caja"]
                line["unidades_por_caja"]      = up_match * pc_match
                line["markup_unidad_pct"]      = p.get("markup_unidad_pct") or line["markup_unidad_pct"]
                line["markup_paquete_pct"]     = p.get("markup_paquete_pct") or line["markup_paquete_pct"]
                line["markup_caja_pct"]        = p.get("markup_caja_pct") or line["markup_caja_pct"]
                line["markup_millar_pct"]      = p.get("markup_millar_pct") or line.get("markup_millar_pct", 40.0)
                line["markup_kg_pct"]          = p.get("markup_kg_pct") or line.get("markup_kg_pct", 40.0)
                line["markup_rollo_pct"]       = p.get("markup_rollo_pct") or line.get("markup_rollo_pct", 40.0)
                line["markup_metro_pct"]       = p.get("markup_metro_pct") or line.get("markup_metro_pct", 40.0)
                line["multiplicador_rollo"]    = p.get("multiplicador_rollo") or line.get("multiplicador_rollo", 1.0)
                line["multiplicador_metro"]    = p.get("multiplicador_metro") or line.get("multiplicador_metro", 1.0)
                line["venta_unidad"]           = p.get("venta_unidad") if p.get("venta_unidad") is not None else line["venta_unidad"]
                line["venta_paquete"]          = p.get("venta_paquete") if p.get("venta_paquete") is not None else line["venta_paquete"]
                line["venta_caja"]             = p.get("venta_caja") if p.get("venta_caja") is not None else line["venta_caja"]
                line["venta_millar"]           = p.get("venta_millar") if p.get("venta_millar") is not None else line.get("venta_millar", False)
                line["venta_kg"]               = p.get("venta_kg") if p.get("venta_kg") is not None else line.get("venta_kg", False)
                line["venta_rollo"]            = p.get("venta_rollo") if p.get("venta_rollo") is not None else line.get("venta_rollo", False)
                line["venta_metro"]            = p.get("venta_metro") if p.get("venta_metro") is not None else line.get("venta_metro", False)
                line["es_materia_prima"]       = bool(p.get("es_materia_prima"))
                line["costo_anterior"]         = float(p.get("costo_unidad_sin_iva") or 0)

                # Aviso no bloqueante: si el emparejamiento fue por codigo de proveedor
                # (no por nombre) y la descripcion de esta linea es muy distinta a la ya
                # guardada, puede ser un codigo reutilizado por el proveedor para otro
                # producto (ver caso HITO). No cambia el match ni el guardado, solo avisa.
                nombre_guardado = p.get("nombre_factura") or p.get("nombre_punto_rojo") or ""
                if match_info.get("via") == "sku" and nombres_muy_distintos(line["nombre_factura"], nombre_guardado):
                    line["aviso_nombre_distinto"] = True
                    line["nombre_guardado_anterior"] = nombre_guardado

    except Exception as e:
        invoice["supabase_match_error"] = str(e)

    invoice["lineas"] = add_calcs(invoice["lineas"], iva_mode)
    invoice["pdf_base64"] = pdf_base64

    return invoice


def recalcular_retefuente_grupo(supabase: Client, proveedor: str, fecha_factura: str, aplica_rete: str) -> None:
    if aplica_rete != "SI":
        return
    try:
        # Mismo fix que en estimar_retefuente: NULL en vigente_desde/vigente_hasta
        # no debe excluir el parámetro vigente.
        params = supabase.table("parametros_retefuente").select(
            "porcentaje,base_minima"
        ).eq("aplica_a", "COMPRAS").eq("activo", True) \
         .or_(f"vigente_desde.is.null,vigente_desde.lte.{fecha_factura}") \
         .or_(f"vigente_hasta.is.null,vigente_hasta.gte.{fecha_factura}").limit(1).execute()
        if not params.data:
            return
        pct_rete = float(params.data[0].get("porcentaje") or 2.5)
        base_min = float(params.data[0].get("base_minima") or 1148000)
        facturas = supabase.table("facturas_contables").select(
            "id,subtotal,retefuente"
        ).eq("proveedor", proveedor).eq("fecha_factura", fecha_factura).execute()
        if not facturas.data:
            return
        total_subtotal = sum(float(f.get("subtotal") or 0) for f in facturas.data)
        if total_subtotal >= base_min:
            for f in facturas.data:
                subtotal_f = float(f.get("subtotal") or 0)
                rete_f = round(subtotal_f * pct_rete / 100, 2)
                valor_pagar_query = supabase.table("facturas_contables").select(
                    "subtotal,iva,valor_descuento"
                ).eq("id", f["id"]).maybe_single().execute()
                if valor_pagar_query.data:
                    vd = valor_pagar_query.data
                    nuevo_valor = max(0.0, float(vd.get("subtotal") or 0) + float(vd.get("iva") or 0) - float(vd.get("valor_descuento") or 0) - rete_f)
                    supabase.table("facturas_contables").update({
                        "retefuente": rete_f,
                        "valor_a_pagar": nuevo_valor
                    }).eq("id", f["id"]).execute()
        else:
            for f in facturas.data:
                if float(f.get("retefuente") or 0) > 0:
                    valor_pagar_query = supabase.table("facturas_contables").select(
                        "subtotal,iva,valor_descuento"
                    ).eq("id", f["id"]).maybe_single().execute()
                    if valor_pagar_query.data:
                        vd = valor_pagar_query.data
                        nuevo_valor = max(0.0, float(vd.get("subtotal") or 0) + float(vd.get("iva") or 0) - float(vd.get("valor_descuento") or 0))
                        supabase.table("facturas_contables").update({
                            "retefuente": 0,
                            "valor_a_pagar": nuevo_valor
                        }).eq("id", f["id"]).execute()
    except Exception:
        pass


def generate_sku(supabase: Client, categoria: str) -> str:
    import unicodedata, re
    prefix = categoria.strip().upper() if categoria else "PROD"
    prefix = ''.join(
        c for c in unicodedata.normalize('NFD', prefix)
        if unicodedata.category(c) != 'Mn'
    )
    prefix = ''.join(c for c in prefix if c.isalnum())[:8]
    if not prefix:
        prefix = "PROD"
    # Usamos el número más alto ya usado (no un conteo), porque si algún producto
    # se elimina queda un "hueco" en la numeración y un conteo simple puede volver
    # a generar un SKU que ya existe (causaba error de duplicado).
    max_num = 0
    try:
        r = supabase.table("productos").select("sku_interno").like("sku_interno", f"{prefix}-%").execute()
        for row in (r.data or []):
            m = re.match(rf"^{re.escape(prefix)}-(\d+)$", row.get("sku_interno") or "")
            if m:
                max_num = max(max_num, int(m.group(1)))
    except Exception:
        pass
    return f"{prefix}-{str(max_num + 1).zfill(4)}"


def registrar_auditoria(sb, accion: str, entidad: str, entidad_id: str = "", descripcion: str = "",
                        usuario_email: str = "", usuario_nombre: str = "", datos_anteriores: dict = None, datos_nuevos: dict = None):
    """Registra una acción en la tabla de auditoría. Falla silenciosamente para no bloquear el flujo principal."""
    try:
        sb.table("auditoria").insert({
            "accion":           accion,
            "entidad":          entidad,
            "entidad_id":       str(entidad_id),
            "descripcion":      descripcion,
            "usuario_email":    usuario_email,
            "usuario_nombre":   usuario_nombre,
            "datos_anteriores": datos_anteriores,
            "datos_nuevos":     datos_nuevos,
        }).execute()
    except Exception:
        pass


@app.post("/save-invoice")
async def save_invoice_endpoint(data: dict):
    try:
        sb = get_supabase()
        invoice  = data.get("invoice", {})
        lineas   = data.get("lineas", [])
        iva_mode = data.get("iva_mode", "NO_INCLUIDO")

        # 1. Buscar o crear proveedor
        nit    = invoice.get("proveedor_nit", "")
        nombre = invoice.get("proveedor", "")
        r = sb.table("proveedores").select("id").eq("nit", nit).limit(1).execute() if nit else None
        if r and r.data:
            proveedor_id = r.data[0]["id"]
        else:
            r2 = sb.table("proveedores").insert({"nombre": nombre, "nit": nit}).execute()
            proveedor_id = r2.data[0]["id"]

        # 1b. Vincular NIT en proveedores_contables — solo si el usuario confirmó
        # explícitamente en el preview que este proveedor (encontrado por nombre
        # parecido, sin NIT propio) es el mismo de esta factura. Antes se hacía
        # automático y en silencio, lo que podía vincular mal proveedores con
        # nombres parecidos pero configuración distinta (ej. dos perfiles del
        # mismo proveedor real para líneas de producto diferentes).
        confirmar_pc_id = data.get("confirmar_proveedor_id")
        if nit and confirmar_pc_id:
            try:
                pc_check = sb.table("proveedores_contables").select("id,nit").eq("id", confirmar_pc_id).maybe_single().execute()
                if pc_check.data and not pc_check.data.get("nit"):
                    sb.table("proveedores_contables").update({"nit": nit}).eq("id", confirmar_pc_id).execute()
            except Exception:
                pass

        # 1c. Subir PDF de la factura a Storage (si viene en el payload) y obtener URL pública
        pdf_url = None
        pdf_b64 = invoice.get("pdf_base64")
        if pdf_b64:
            try:
                import base64 as b64lib
                pdf_bytes = b64lib.b64decode(pdf_b64)
                numero_fact_safe = re.sub(r"[^A-Za-z0-9_-]", "_", str(invoice.get("numero_factura") or "factura"))
                pdf_path = f"{numero_fact_safe}_{int(datetime.now().timestamp())}.pdf"
                sb.storage.from_("facturas-pdf").upload(
                    pdf_path, pdf_bytes, {"content-type": "application/pdf"}
                )
                pdf_url = sb.storage.from_("facturas-pdf").get_public_url(pdf_path)
            except Exception as pdf_err:
                print(f"ERROR subiendo PDF a Storage: {pdf_err}")
                pdf_url = None
        else:
            print("No vino pdf_base64 en el payload de save-invoice")

        # 2. Insertar factura
        fac = sb.table("facturas").insert({
            "proveedor_id":   proveedor_id,
            "proveedor_nit":  nit,
            "numero_factura": invoice.get("numero_factura"),
            "fecha":          invoice.get("fecha"),
            "subtotal":       invoice.get("subtotal_factura"),
            "iva":            invoice.get("iva_factura"),
            "total":          invoice.get("total_factura"),
            "iva_modo":       iva_mode,
            "cufe":           invoice.get("cufe"),
            "archivo":        invoice.get("archivo_nombre", ""),
            "nota":           data.get("nota_factura") or None,
            "pdf_url":        pdf_url,
        }).execute()
        factura_id = fac.data[0]["id"]

        # 3. Insertar/actualizar productos e historial
        for line in lineas:
            prod_id = line.get("producto_id")

            # Red de seguridad: si por algún motivo el match no llegó desde el preview
            # (ej. datos editados manualmente), reintentar la búsqueda por nombre antes
            # de insertar, para no chocar contra la unique constraint.
            if not prod_id and not line.get("nombre_duplicado_sin_sku"):
                raw_sku_linea = (line.get("sku_proveedor") or "").strip()
                if not raw_sku_linea or _es_sku_temporal(raw_sku_linea):
                    existente = sb.table("productos").select("id").eq(
                        "proveedor_id", proveedor_id
                    ).eq("sku_proveedor", "").eq(
                        "nombre_factura", line.get("nombre_factura", "")
                    ).eq("activo", True).limit(1).execute()
                    if existente.data:
                        prod_id = existente.data[0]["id"]

            up      = int(line.get("unidades_por_paquete") or 1)
            pc      = int(line.get("paquetes_por_caja") or 1)
            uc      = up * pc

            precio_fact      = float(line.get("precio_unitario_factura") or 0)
            desc_pct_l       = float(line.get("descuento_factura_pct") or 0)
            precio_fact_base = precio_fact  # guardar original antes del descuento
            if line.get("descuento_afecta_costo") and desc_pct_l > 0:
                precio_fact = money(precio_fact * (1 - desc_pct_l / 100))

            iva_mode_s  = data.get("iva_mode", "NO_INCLUIDO")
            iva_pct_l   = float(line.get("iva_porcentaje") or IVA_DEFAULT)
            costo_base  = cost_without_tax(precio_fact, iva_mode_s, iva_pct_l)
            transporte  = float(line.get("transporte_adicional") or 0)
            costo_final = costo_base + transporte
            pres_s      = line.get("presentacion_facturada", "Unidad")

            # Usar precio_es_por del usuario si viene, si no inferir
            precio_es_por_s = line.get("precio_es_por") or ""
            if not precio_es_por_s:
                precio_es_por_s = inferir_precio_es_por(costo_final, up, pc, uc)

            upm_s = int(line.get("unidades_por_millar") or 1000)
            cu, cp, cc = calc_costs(costo_final, pres_s, up, pc, precio_es_por_s, upm_s)

            if prod_id:
                old = sb.table("productos").select("costo_unidad_sin_iva").eq("id", prod_id).maybe_single().execute()
                costo_ant = float(old.data.get("costo_unidad_sin_iva") or 0) if old.data else 0
                variacion = round(((cu - costo_ant) / costo_ant * 100), 2) if costo_ant > 0 else 0

                sb.table("productos").update({
                    "costo_unidad_sin_iva":   cu,
                    "costo_paquete_sin_iva":  cp,
                    "costo_caja_sin_iva":     cc,
                    "unidades_por_paquete":   up,
                    "paquetes_por_caja":      pc,
                    "unidades_por_caja":      uc,
                    "ultima_factura":         invoice.get("numero_factura"),
                    "ultima_fecha":           invoice.get("fecha"),
                    "markup_unidad_pct":      float(line.get("markup_unidad_pct") or 40),
                    "markup_paquete_pct":     float(line.get("markup_paquete_pct") or 35),
                    "markup_caja_pct":        float(line.get("markup_caja_pct") or 31.58),
                    "markup_millar_pct":      float(line.get("markup_millar_pct") or 40),
                    "markup_kg_pct":          float(line.get("markup_kg_pct") or 40),
                    "markup_rollo_pct":       float(line.get("markup_rollo_pct") or 40),
                    "markup_metro_pct":       float(line.get("markup_metro_pct") or 40),
                    "multiplicador_rollo":    float(line.get("multiplicador_rollo") or 1),
                    "multiplicador_metro":    float(line.get("multiplicador_metro") or 1),
                    **({"es_materia_prima": True} if line.get("es_materia_prima") else {}),
                    "venta_unidad":           bool(line.get("venta_unidad")),
                    "venta_paquete":          bool(line.get("venta_paquete")),
                    "venta_caja":             bool(line.get("venta_caja")),
                    "venta_millar":           bool(line.get("venta_millar")),
                    "venta_kg":               bool(line.get("venta_kg")),
                    "venta_rollo":            bool(line.get("venta_rollo")),
                    "venta_metro":            bool(line.get("venta_metro")),
                    "nota_descuento":         line.get("nota_descuento") or "",
                    "precio_factura_base":    precio_fact_base,
                    "precio_es_por":          precio_es_por_s,
                    "unidades_por_millar":    upm_s,
                    "iva_porcentaje":         iva_pct_l,
                }).eq("id", prod_id).execute()

                estado = "NUEVO" if costo_ant == 0 else ("SUBIO" if cu > costo_ant else "BAJO" if cu < costo_ant else "SIN_CAMBIO")
            else:
                costo_ant = 0
                variacion = 0
                estado    = "NUEVO"
                categoria_line = line.get("categoria", "") or ""
                sku_int = generate_sku(sb, categoria_line)

                # Limpiar SKU temporal antes de guardar.
                # Un SKU tipo "L001" lo asignamos nosotros cuando el proveedor no manda
                # código real. Si lo guardamos en la BD, futuras facturas del mismo
                # proveedor harán match falso contra productos completamente distintos.
                sku_guardar = line.get("sku_proveedor", "") or ""
                # Si el nombre se repite en otra línea de ESTA factura sin SKU real, no
                # limpiar a "" — chocaría contra la otra línea en la unique constraint y
                # fusionaría dos productos que en realidad son distintos.
                if _es_sku_temporal(sku_guardar) and not line.get("nombre_duplicado_sin_sku"):
                    sku_guardar = ""

                producto_payload = {
                    "proveedor_id":           proveedor_id,
                    "sku_proveedor":          sku_guardar,
                    "sku_interno":            sku_int,
                    "nombre_factura":         line.get("nombre_factura", ""),
                    "nombre_punto_rojo":      line.get("nombre_punto_rojo") or line.get("nombre_factura", ""),
                    "categoria":              line.get("categoria", ""),
                    "subcategoria":           line.get("subcategoria", "") or "",
                    "presentacion_facturada": line.get("presentacion_facturada", "Unidad"),
                    "unidades_por_paquete":   up,
                    "paquetes_por_caja":      pc,
                    "unidades_por_caja":      uc,
                    "costo_unidad_sin_iva":   cu,
                    "costo_paquete_sin_iva":  cp,
                    "costo_caja_sin_iva":     cc,
                    "markup_unidad_pct":      float(line.get("markup_unidad_pct") or 40),
                    "markup_paquete_pct":     float(line.get("markup_paquete_pct") or 35),
                    "markup_caja_pct":        float(line.get("markup_caja_pct") or 31.58),
                    "markup_millar_pct":      float(line.get("markup_millar_pct") or 40),
                    "markup_kg_pct":          float(line.get("markup_kg_pct") or 40),
                    "markup_rollo_pct":       float(line.get("markup_rollo_pct") or 40),
                    "markup_metro_pct":       float(line.get("markup_metro_pct") or 40),
                    "multiplicador_rollo":    float(line.get("multiplicador_rollo") or 1),
                    "multiplicador_metro":    float(line.get("multiplicador_metro") or 1),
                    "venta_unidad":           bool(line.get("venta_unidad")),
                    "venta_paquete":          bool(line.get("venta_paquete")),
                    "venta_caja":             bool(line.get("venta_caja")),
                    "venta_millar":           bool(line.get("venta_millar")),
                    "venta_kg":               bool(line.get("venta_kg")),
                    "venta_rollo":            bool(line.get("venta_rollo")),
                    "venta_metro":            bool(line.get("venta_metro")),
                    "es_materia_prima":       bool(line.get("es_materia_prima")),
                    "activo":                 True,
                    "ultima_factura":         invoice.get("numero_factura"),
                    "ultima_fecha":           invoice.get("fecha"),
                    "nota_descuento":         line.get("nota_descuento") or "",
                    "precio_factura_base":    precio_fact_base,
                    "precio_es_por":          precio_es_por_s,
                    "unidades_por_millar":    upm_s,
                    "iva_porcentaje":         iva_pct_l,
                }

                # Si el SKU es real (no vacío), usar upsert con conflict para actualizar
                # si ya existe ese SKU+proveedor+nombre. Si el SKU es vacío (proveedor sin
                # código), usar insert puro para evitar colisiones entre productos distintos
                # del mismo proveedor que comparten sku_proveedor="" y nombre_factura similar.
                if sku_guardar:
                    res = sb.table("productos").upsert(
                        producto_payload,
                        on_conflict="proveedor_id,sku_proveedor,nombre_factura"
                    ).execute()
                else:
                    res = sb.table("productos").insert(producto_payload).execute()

                prod_id = res.data[0]["id"]

            # Historial de costos
            sb.table("historial_costos").insert({
                "factura_id":                   factura_id,
                "producto_id":                  prod_id,
                "estado":                       estado,
                "presentacion_facturada":       line.get("presentacion_facturada"),
                "costo_presentacion_facturada": float(line.get("costo_caja_sin_iva") or cu),
                "costo_unidad_anterior":        costo_ant,
                "costo_unidad_nuevo":           cu,
                "variacion_porcentaje":         variacion,
                "venta_unidad":                 bool(line.get("venta_unidad")),
                "venta_paquete":                bool(line.get("venta_paquete")),
                "venta_caja":                   bool(line.get("venta_caja")),
                "venta_millar":                 bool(line.get("venta_millar")),
                "venta_kg":                     bool(line.get("venta_kg")),
                "venta_rollo":                  bool(line.get("venta_rollo")),
                "venta_metro":                  bool(line.get("venta_metro")),
                "markup_unidad_pct":            float(line.get("markup_unidad_pct") or 40),
                "markup_paquete_pct":           float(line.get("markup_paquete_pct") or 35),
                "markup_caja_pct":              float(line.get("markup_caja_pct") or 31.58),
                "markup_millar_pct":            float(line.get("markup_millar_pct") or 40),
                "markup_kg_pct":                float(line.get("markup_kg_pct") or 40),
                "markup_rollo_pct":             float(line.get("markup_rollo_pct") or 40),
                "markup_metro_pct":             float(line.get("markup_metro_pct") or 40),
                "precio_factura_original":      precio_fact_base,
                "precio_es_por":                precio_es_por_s,
                "nota":                         (line.get("nota") or "").strip() or None,
            }).execute()

            # Historial de descuentos
            desc_pct_hist = float(line.get("descuento_factura_pct") or 0)
            if desc_pct_hist > 0:
                desc_aplicado = bool(line.get("descuento_afecta_costo", False))
                nota_hist = line.get("nota_descuento") or f"Descuento {desc_pct_hist:.0f}% en factura"
                if desc_aplicado:
                    nota_hist = f"{nota_hist} — aplicado al costo"
                else:
                    nota_hist = f"{nota_hist} — NO aplicado al costo"
                sb.table("historial_descuentos").insert({
                    "producto_id":   prod_id,
                    "descuento_pct": desc_pct_hist,
                    "nota":          nota_hist,
                    "factura_id":    factura_id,
                }).execute()
                sb.table("productos").update({
                    "descuento_pct_factura": desc_pct_hist,
                    "descuento_aplicado":    desc_aplicado,
                }).eq("id", prod_id).execute()

        # 4. Factura contable
        prov_info  = data.get("proveedor_info", {})

        # Forma de pago: priorizar XML, luego proveedores_contables, default CREDITO
        forma_pago_xml = invoice.get("forma_pago_xml", "")
        prov_forma_pago = prov_info.get("forma_pago") or ""
        # Si el proveedor ya tiene forma de pago configurada, usarla
        # Solo usar el XML si el proveedor es nuevo o no tiene forma de pago
        if prov_forma_pago:
            forma_pago = prov_forma_pago
        else:
            forma_pago = forma_pago_xml or "CREDITO"

        # Actualizar forma_pago y requiere_acuse en proveedores_contables si tenemos id
        # Regla de negocio: el acuse DIAN solo aplica a facturas de CREDITO (no CONTADO, no OTRO)
        pc_id = prov_info.get("id")
        if pc_id and forma_pago_xml:
            try:
                sb.table("proveedores_contables").update({
                    "forma_pago": forma_pago,
                }).eq("id", pc_id).execute()
            except Exception:
                pass

        # Si proveedor no existe en proveedores_contables, crearlo
        if not pc_id and nombre:
            try:
                nuevo_pc = sb.table("proveedores_contables").insert({
                    "proveedor_nombre": nombre,
                    "nit": nit or "",
                    "forma_pago": forma_pago,
                    "aplica_retefuente": "NO",
                    "descuento_pct": 0,
                    "regimen": "COMUN",
                    "descuento_afecta_costo": False,
                }).execute()
                if nuevo_pc.data:
                    pc_id = nuevo_pc.data[0]["id"]
            except Exception as e_insert:
                # El insert puede fallar por conflict de NIT duplicado u otro error.
                # Intentamos buscar si ya existe con ese NIT o nombre antes de rendirse.
                try:
                    busqueda = None
                    if nit:
                        busqueda = sb.table("proveedores_contables").select("id").eq("nit", nit).maybe_single().execute()
                    if not busqueda or not busqueda.data:
                        busqueda = sb.table("proveedores_contables").select("id").ilike("proveedor_nombre", nombre).maybe_single().execute()
                    if busqueda and busqueda.data:
                        pc_id = busqueda.data["id"]
                    else:
                        # Último intento: upsert por NIT
                        raise Exception(f"No se pudo crear proveedor contable: {str(e_insert)}")
                except Exception:
                    pass  # El proveedor no quedó en contables — se detectará en la próxima verificación

        # Verificación final: si después de todo el proveedor no existe en contables, crearlo
        if not pc_id and nombre:
            try:
                verificar = sb.table("proveedores_contables").select("id").ilike("proveedor_nombre", nombre).maybe_single().execute()
                if not verificar or not verificar.data:
                    fallback = sb.table("proveedores_contables").insert({
                        "proveedor_nombre": nombre,
                        "nit": nit or "",
                        "forma_pago": "CONTADO",
                        "aplica_retefuente": "NO",
                        "descuento_pct": 0,
                        "regimen": "COMUN",
                        "descuento_afecta_costo": False,
                    }).execute()
                    if fallback.data:
                        pc_id = fallback.data[0]["id"]
            except Exception:
                pass
        desc_pct   = float(prov_info.get("descuento_pct") or 0)
        subtotal   = float(invoice.get("subtotal_factura") or 0)
        iva_val    = float(invoice.get("iva_factura") or 0)
        valor_desc = round(subtotal * desc_pct / 100, 2)

        retefuente_xml = float(invoice.get("retefuente_xml") or 0)
        aplica_rete    = prov_info.get("aplica_retefuente", "NO")
        retefuente     = estimar_retefuente(sb, prov_info, subtotal, invoice.get("fecha") or "", retefuente_xml)

        valor_pagar = max(0.0, subtotal + iva_val - valor_desc - retefuente)

        marcar_acuse = bool(data.get("marcar_acuse", False))

        sb.table("facturas_contables").insert({
            "factura_id":        factura_id,
            "proveedor":         nombre,
            "numero_factura":    invoice.get("numero_factura"),
            "fecha_factura":     invoice.get("fecha"),
            "fecha_revision":    str(date.today()),
            "forma_pago":        forma_pago,
            "subtotal":          subtotal,
            "descuento_pct":     desc_pct,
            "valor_descuento":   valor_desc,
            "aplica_retefuente": aplica_rete,
            "retefuente":        retefuente,
            "iva":               iva_val,
            "valor_a_pagar":     valor_pagar,
            "precios_revisados": "NO",
            "cufe":              invoice.get("cufe", ""),
            "acuse":             "SI" if marcar_acuse else "NO",
            "fecha_acuse":       str(date.today()) if marcar_acuse else None,
        }).execute()

        if not retefuente_xml:
            recalcular_retefuente_grupo(sb, nombre, invoice.get("fecha", ""), aplica_rete)

        # Registrar auditoría
        usuario_email  = data.get("usuario_email", "")
        usuario_nombre = data.get("usuario_nombre", "")
        registrar_auditoria(sb,
            accion="GUARDAR_FACTURA", entidad="facturas", entidad_id=str(factura_id),
            descripcion=f"Factura {invoice.get('numero_factura')} de {nombre} — ${subtotal:,.0f}",
            usuario_email=usuario_email, usuario_nombre=usuario_nombre,
        )

        return {"ok": True, "factura_id": factura_id, "mensaje": f"Factura {invoice.get('numero_factura')} guardada correctamente."}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/parse-credit-note")
async def parse_credit_note_endpoint(file: UploadFile = File(...)):
    raw = await file.read()
    if len(raw) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="Archivo demasiado grande. Máximo 10 MB.")
    filename = file.filename or ""

    try:
        if filename.lower().endswith(".zip"):
            with zipfile.ZipFile(io.BytesIO(raw)) as z:
                xml_names = [n for n in z.namelist() if n.lower().endswith(".xml") and not n.lower().startswith("__")]
                if not xml_names:
                    raise HTTPException(status_code=400, detail="El ZIP no contiene XML.")
                raw_xml = z.read(xml_names[0])
        elif filename.lower().endswith(".xml"):
            raw_xml = raw
        else:
            raise HTTPException(status_code=400, detail="Solo se aceptan XML o ZIP.")

        root_el = extract_invoice_xml(raw_xml)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Error al leer XML: {str(e)}")

    try:
        supplier = (
            first_text(root_el, ["AccountingSupplierParty","Party","PartyLegalEntity","RegistrationName"]) or
            first_text(root_el, ["AccountingSupplierParty","Party","PartyName","Name"]) or ""
        )
        nit = (
            first_text(root_el, ["AccountingSupplierParty","Party","PartyTaxScheme","CompanyID"]) or
            first_text(root_el, ["AccountingSupplierParty","Party","PartyLegalEntity","CompanyID"]) or ""
        )
        numero_nota = first_text(root_el, ["ID"]) or ""
        fecha_nota  = first_text(root_el, ["IssueDate"]) or str(date.today())

        factura_original = ""
        for ref in all_descendants(root_el, "BillingReference"):
            factura_original = first_text(ref, ["InvoiceDocumentReference","ID"]) or ""
            if factura_original:
                break
        if not factura_original:
            for dr in all_descendants(root_el, "DiscrepancyResponse"):
                factura_original = first_text(dr, ["ReferenceID"]) or ""
                if factura_original:
                    break

        subtotal = parse_decimal(first_text(root_el, ["LegalMonetaryTotal","LineExtensionAmount"]))
        total    = parse_decimal(first_text(root_el, ["LegalMonetaryTotal","PayableAmount"]))
        iva      = money(total - subtotal) if subtotal and total and total >= subtotal else 0.0

        retefuente = 0.0
        for ta in all_descendants(root_el, "TaxTotal"):
            tax_id   = ""
            tax_name = ""
            for sc in all_descendants(ta, "TaxScheme"):
                tax_id   = first_text(sc, ["ID"]) or ""
                tax_name = (first_text(sc, ["Name"]) or "").upper()
            if tax_id in ("06","05","07") or "RETE" in tax_name:
                retefuente += parse_decimal(first_text(ta, ["TaxAmount"]))

        motivo = ""
        for dr in all_descendants(root_el, "DiscrepancyResponse"):
            motivo = first_text(dr, ["Description"]) or ""
            if motivo:
                break
        if not motivo:
            for note in all_descendants(root_el, "Note"):
                if note.text and note.text.strip():
                    motivo = note.text.strip()
                    break

        factura_contable_id = None
        factura_contable_advertencia = None
        try:
            sb = get_supabase()
            if factura_original:
                r = sb.table("facturas_contables").select("id,subtotal,iva,retefuente,valor_a_pagar,proveedor").eq("numero_factura", factura_original).limit(1).execute()
                if r.data:
                    factura_contable_id = r.data[0]["id"]
                else:
                    factura_contable_advertencia = f"No se encontró la factura '{factura_original}' en contabilidad. Verifica el número antes de guardar."
        except Exception as e:
            factura_contable_advertencia = f"Error al buscar factura contable: {str(e)}"

        return {
            "proveedor":                    supplier,
            "proveedor_nit":                nit,
            "numero_nota":                  numero_nota,
            "fecha_nota":                   fecha_nota,
            "factura_original":             factura_original,
            "factura_contable_id":          factura_contable_id,
            "factura_contable_advertencia": factura_contable_advertencia,
            "subtotal":                     money(subtotal),
            "iva":                          money(iva),
            "retefuente":                   money(retefuente),
            "total":                        money(total),
            "motivo":                       motivo,
        }

    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Error al parsear nota crédito: {str(e)}")


@app.post("/toggle-descuento")
async def toggle_descuento(data: dict):
    try:
        sb       = get_supabase()
        prod_id  = data.get("producto_id")
        aplicado = bool(data.get("descuento_aplicado", False))
        nota     = data.get("nota", "")

        prod = sb.table("productos").select(
            "precio_factura_base,precio_es_por,descuento_pct_factura,"
            "unidades_por_paquete,paquetes_por_caja,presentacion_facturada,iva_porcentaje,unidades_por_millar"
        ).eq("id", prod_id).maybe_single().execute()

        if not prod.data:
            raise HTTPException(status_code=404, detail="Producto no encontrado")

        p             = prod.data
        precio_base   = float(p.get("precio_factura_base") or 0)
        desc_pct      = float(p.get("descuento_pct_factura") or 0)
        precio_es_por = p.get("precio_es_por") or ""
        pres          = p.get("presentacion_facturada") or "Unidad"
        up            = int(p.get("unidades_por_paquete") or 1)
        pc            = int(p.get("paquetes_por_caja") or 1)
        upm           = int(p.get("unidades_por_millar") or 1000)
        iva_pct       = float(p.get("iva_porcentaje") or IVA_DEFAULT)

        if precio_base <= 0:
            raise HTTPException(status_code=400, detail="Este producto no tiene precio base guardado. Sube la factura nuevamente para activar esta función.")

        precio_fact = precio_base
        if aplicado and desc_pct > 0:
            precio_fact = money(precio_base * (1 - desc_pct / 100))

        costo_base = cost_without_tax(precio_fact, "NO_INCLUIDO", iva_pct)
        cu, cp, cc = calc_costs(costo_base, pres, up, pc, precio_es_por, upm)

        sb.table("productos").update({
            "costo_unidad_sin_iva":  cu,
            "costo_paquete_sin_iva": cp,
            "costo_caja_sin_iva":    cc,
            "descuento_aplicado":    aplicado,
            "nota_descuento":        nota,
        }).eq("id", prod_id).execute()

        sb.table("historial_descuentos").insert({
            "producto_id":   prod_id,
            "descuento_pct": desc_pct if aplicado else 0,
            "nota":          nota or ("Descuento activado desde panel" if aplicado else "Descuento desactivado desde panel"),
        }).execute()

        return {
            "ok": True,
            "costo_unidad_sin_iva":  cu,
            "costo_paquete_sin_iva": cp,
            "costo_caja_sin_iva":    cc,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/recalc-precio-es-por")
async def recalc_precio_es_por(data: dict):
    """Recalcula los costos desde el precio original de factura usando el nuevo precio_es_por."""
    try:
        sb            = get_supabase()
        prod_id       = data.get("producto_id")
        precio_es_por = data.get("precio_es_por", "")
        pres          = data.get("presentacion_facturada", "")
        up            = int(data.get("unidades_por_paquete") or 1)
        pc            = int(data.get("paquetes_por_caja") or 1)
        upm           = int(data.get("unidades_por_millar") or 0)

        prod = sb.table("productos").select(
            "precio_factura_base,descuento_pct_factura,descuento_aplicado,"
            "presentacion_facturada,unidades_por_paquete,paquetes_por_caja,iva_porcentaje,unidades_por_millar"
        ).eq("id", prod_id).maybe_single().execute()

        if not prod.data:
            raise HTTPException(status_code=404, detail="Producto no encontrado")

        p            = prod.data
        precio_base  = float(p.get("precio_factura_base") or 0)

        if precio_base <= 0:
            raise HTTPException(status_code=400, detail="Este producto no tiene precio base guardado. Sube la factura nuevamente.")

        desc_pct      = float(p.get("descuento_pct_factura") or 0)
        desc_aplicado = bool(p.get("descuento_aplicado") or False)
        iva_pct       = float(p.get("iva_porcentaje") or IVA_DEFAULT)

        pres_final = pres or p.get("presentacion_facturada") or "Unidad"
        up_final   = up or int(p.get("unidades_por_paquete") or 1)
        pc_final   = pc or int(p.get("paquetes_por_caja") or 1)
        upm_final  = upm or int(p.get("unidades_por_millar") or 1000)
        uc_final   = up_final * pc_final

        precio_fact = precio_base
        if desc_aplicado and desc_pct > 0:
            precio_fact = money(precio_base * (1 - desc_pct / 100))

        costo_base_val = cost_without_tax(precio_fact, "NO_INCLUIDO", iva_pct)

        # Si no viene precio_es_por, inferir automáticamente
        if not precio_es_por:
            precio_es_por = inferir_precio_es_por(costo_base_val, up_final, pc_final, uc_final)

        cu, cp, cc = calc_costs(costo_base_val, pres_final, up_final, pc_final, precio_es_por, upm_final)

        sb.table("productos").update({
            "precio_es_por":          precio_es_por,
            "presentacion_facturada": pres_final,
            "unidades_por_paquete":   up_final,
            "paquetes_por_caja":      pc_final,
            "unidades_por_caja":      uc_final,
            "unidades_por_millar":    upm_final,
            "costo_unidad_sin_iva":   cu,
            "costo_paquete_sin_iva":  cp,
            "costo_caja_sin_iva":     cc,
        }).eq("id", prod_id).execute()

        return {
            "ok": True,
            "costo_unidad_sin_iva":  cu,
            "costo_paquete_sin_iva": cp,
            "costo_caja_sin_iva":    cc,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── Materias primas y productos derivados ───────────────────────────────────

@app.get("/materias-primas")
async def listar_materias_primas():
    """Lista productos marcados como materia prima, para el selector de 'crear derivado de...'."""
    try:
        sb = get_supabase()
        r = sb.table("productos").select(
            "id,nombre_punto_rojo,sku_interno,categoria,costo_unidad_sin_iva,precio_es_por,unidades_por_paquete,paquetes_por_caja"
        ).eq("es_materia_prima", True).eq("activo", True).order("nombre_punto_rojo").execute()
        return {"ok": True, "items": r.data or []}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/categorias")
async def listar_categorias():
    """Lista categorías únicas ordenadas por categorias_orden.orden, luego alfabético."""
    try:
        sb = get_supabase()
        # Obtener categorías de productos
        r = sb.table("productos").select("categoria").eq("activo", True).execute()
        cats_productos = set()
        for row in (r.data or []):
            cat = (row.get("categoria") or "").strip()
            if cat:
                cats_productos.add(cat)

        # Obtener orden guardado
        orden_r = sb.table("categorias_orden").select("nombre,orden").execute()
        orden_map = {row["nombre"]: row["orden"] for row in (orden_r.data or [])}

        # Registrar categorías nuevas que no están en categorias_orden -- insertadas
        # en su posición alfabética dentro del orden actual, no siempre al final,
        # igual que ya hace el botón "Crear categoría" en Admin.
        nuevas = [c for c in cats_productos if c not in orden_map]
        if nuevas:
            lista_actual = [n for n, _ in sorted(orden_map.items(), key=lambda kv: kv[1])]
            for nueva in sorted(nuevas):
                pos = len(lista_actual)
                for idx, nombre_existente in enumerate(lista_actual):
                    if nombre_existente.upper() > nueva.upper():
                        pos = idx
                        break
                lista_actual.insert(pos, nueva)
            try:
                for i, nombre in enumerate(lista_actual):
                    nuevo_orden = (i + 1) * 10
                    if orden_map.get(nombre) != nuevo_orden:
                        sb.table("categorias_orden").upsert({"nombre": nombre, "orden": nuevo_orden}, on_conflict="nombre").execute()
                    orden_map[nombre] = nuevo_orden
            except Exception:
                pass

        # Union: categorías con al menos un producto activo + categorías creadas
        # manualmente desde Admin aunque todavía no tengan ningún producto -- si no,
        # "Crear categoría" guarda bien pero nunca aparece hasta que algo la use.
        todas_categorias = cats_productos | set(orden_map.keys())

        # Ordenar por orden guardado, luego alfabético para empates
        categorias = sorted(todas_categorias, key=lambda c: (orden_map.get(c, 9999), c.upper()))
        return {"ok": True, "categorias": categorias}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/subcategorias")
async def listar_subcategorias(categoria: str = ""):
    """Lista subcategorías únicas de una categoría específica."""
    try:
        sb = get_supabase()
        q = sb.table("productos").select("subcategoria").eq("activo", True)
        if categoria:
            q = q.eq("categoria", categoria)
        r = q.execute()
        vistas = set()
        subcategorias = []
        for row in (r.data or []):
            sub = (row.get("subcategoria") or "").strip()
            if sub and sub.upper() not in vistas:
                vistas.add(sub.upper())
                subcategorias.append(sub)
        subcategorias.sort(key=lambda s: s.upper())
        return {"ok": True, "subcategorias": subcategorias}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/todas-subcategorias")
async def listar_todas_subcategorias():
    """Devuelve todas las subcategorías agrupadas por categoría en una sola query."""
    try:
        sb = get_supabase()
        r = sb.table("productos").select("categoria,subcategoria").eq("activo", True).execute()
        mapa: dict = {}
        for row in (r.data or []):
            cat = (row.get("categoria") or "").strip()
            sub = (row.get("subcategoria") or "").strip()
            if cat and sub:
                if cat not in mapa:
                    mapa[cat] = set()
                mapa[cat].add(sub)
        # Convertir sets a listas ordenadas
        resultado = {cat: sorted(list(subs)) for cat, subs in mapa.items()}
        return {"ok": True, "subcategorias": resultado}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/categorias-orden")
async def guardar_orden_categorias(data: dict):
    """Guarda el orden de categorías después del drag & drop. Recibe {orden: ['CAT1','CAT2',...]}"""
    try:
        sb = get_supabase()
        categorias_ordenadas: list = data.get("orden", [])
        for i, nombre in enumerate(categorias_ordenadas):
            sb.table("categorias_orden").upsert(
                {"nombre": nombre, "orden": (i + 1) * 10},
                on_conflict="nombre"
            ).execute()
        return {"ok": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/subir-imagen-producto")
async def subir_imagen_producto(
    sku_interno: str,
    file: UploadFile = File(...),
):
    """Sube la imagen de un producto a Storage usando la service key (evita
    problemas de RLS/sesión en el navegador) y devuelve la URL pública.

    Siempre se comprime y convierte a JPG con un nombre fijo (principal.jpg)
    para que el upsert reemplace correctamente la imagen anterior en vez de
    crear archivos duplicados con distinta extensión (bug detectado: subir
    una imagen como .png y luego como .jpg generaba dos archivos separados,
    desperdiciando espacio en Supabase Storage)."""
    try:
        import PIL.Image
        sb = get_supabase()
        raw = await file.read()

        # Abrir con PIL, convertir a RGB (por si viene PNG con transparencia) y redimensionar
        img = PIL.Image.open(io.BytesIO(raw))
        if img.mode in ("RGBA", "P", "LA"):
            fondo = PIL.Image.new("RGB", img.size, (255, 255, 255))
            img = img.convert("RGBA")
            fondo.paste(img, mask=img.split()[-1])
            img = fondo
        elif img.mode != "RGB":
            img = img.convert("RGB")

        MAX_ANCHO = 1200
        if img.width > MAX_ANCHO:
            ratio = MAX_ANCHO / img.width
            nuevo_alto = int(img.height * ratio)
            img = img.resize((MAX_ANCHO, nuevo_alto), PIL.Image.LANCZOS)

        buffer = io.BytesIO()
        img.save(buffer, format="JPEG", quality=85, optimize=True)
        data_final = buffer.getvalue()

        sku_safe = re.sub(r"[^A-Za-z0-9_-]", "_", sku_interno)
        path = f"{sku_safe}/principal.jpg"

        # Eliminar cualquier versión previa con otra extensión (residuos de subidas anteriores al fix)
        for ext_vieja in ("png", "webp", "gif", "jpeg"):
            try:
                sb.storage.from_("Productos").remove([f"{sku_safe}/principal.{ext_vieja}"])
            except Exception:
                pass

        sb.storage.from_("Productos").upload(
            path, data_final, {"content-type": "image/jpeg", "upsert": "true"}
        )
        url = sb.storage.from_("Productos").get_public_url(path)
        return {"ok": True, "imagen_url": url}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/crear-producto-derivado")
async def crear_producto_derivado(data: dict):
    """
    Crea un producto derivado de una materia prima (ej: bolsas calculadas a partir
    de un rollo de plástico comprado por kg), usando la fórmula:
      costo_x100 = (ancho * largo * constante * calibre * precio_kg) / 10
    """
    try:
        sb = get_supabase()

        nombre        = (data.get("nombre") or "").strip()
        categoria     = (data.get("categoria") or "").strip()
        subcategoria  = (data.get("subcategoria") or "").strip()
        if not nombre:
            raise HTTPException(status_code=400, detail="El nombre es obligatorio")

        materia_prima_id = data.get("materia_prima_id")

        # Calculadora de bolsas
        ancho     = float(data.get("ancho") or 0)
        largo     = float(data.get("largo") or 0)
        calibre   = float(data.get("calibre") or 0)
        precio_kg = float(data.get("precio_kg") or 0)
        constante = float(data.get("constante") or 0.0302)
        sellado   = float(data.get("sellado") or 0)
        uds_paquete = int(data.get("uds_paquete") or 100)

        # Si se usa la formula, guardamos tambien sus datos (no solo el resultado)
        # para que el trigger de Supabase pueda recalcular este producto cuando
        # cambie el precio de la materia prima.
        usa_formula = ancho > 0 and largo > 0 and calibre > 0 and precio_kg > 0
        costo_unidad = 0.0
        if usa_formula:
            costo_100 = (ancho * largo * constante * calibre * precio_kg) / 10
            costo_unidad = money(costo_100 / 100) + sellado
        else:
            costo_unidad = float(data.get("costo_unidad") or 0)

        up = max(int(data.get("up") or uds_paquete or 1), 1)
        pc = max(int(data.get("pc") or 1), 1)
        uc = up * pc

        cp = money(costo_unidad * up)
        cc = money(costo_unidad * uc)

        sku_int = generate_sku(sb, categoria)

        # presentaciones_extra: lista de presentaciones de venta adicionales (x10, x50, etc.)
        presentaciones_extra = data.get("presentaciones_extra") or []

        # Copiar la ultima factura/fecha de la materia prima base -- sin esto el
        # derivado queda con ultima_factura vacio para siempre, dando la impresion
        # de que nunca quedo ligado a ninguna factura.
        ultima_factura_val = None
        ultima_fecha_val = None
        if materia_prima_id:
            try:
                mp_row = sb.table("productos").select("ultima_factura,ultima_fecha").eq("id", materia_prima_id).maybe_single().execute()
                if mp_row.data:
                    ultima_factura_val = mp_row.data.get("ultima_factura")
                    ultima_fecha_val = mp_row.data.get("ultima_fecha")
            except Exception:
                pass

        insert_payload = {
            "sku_interno":             sku_int,
            "sku_proveedor":           "",
            "nombre_factura":          nombre,
            "nombre_punto_rojo":       nombre,
            "categoria":               categoria,
            "subcategoria":            subcategoria,
            "presentacion_facturada":  data.get("presentacion") or "Unidad",
            "precio_es_por":           "",
            "unidades_por_paquete":    up,
            "paquetes_por_caja":       pc,
            "unidades_por_caja":       uc,
            "costo_unidad_sin_iva":    costo_unidad,
            "costo_paquete_sin_iva":   cp,
            "costo_caja_sin_iva":      cc,
            "markup_unidad_pct":       float(data.get("markup_u") or 40),
            "markup_paquete_pct":      float(data.get("markup_p") or 35),
            "markup_caja_pct":         float(data.get("markup_c") or 31.58),
            "markup_millar_pct":       float(data.get("markup_millar") or 40),
            "markup_kg_pct":           float(data.get("markup_kg") or 40),
            "markup_rollo_pct":        float(data.get("markup_rollo") or 40),
            "markup_metro_pct":        float(data.get("markup_metro") or 40),
            "multiplicador_rollo":     float(data.get("multiplicador_rollo") or 1),
            "multiplicador_metro":     float(data.get("multiplicador_metro") or 1),
            "venta_unidad":            bool(data.get("venta_unidad", True)),
            "venta_paquete":           bool(data.get("venta_paquete", False)),
            "venta_caja":              bool(data.get("venta_caja", False)),
            "venta_millar":            bool(data.get("venta_millar", False)),
            "venta_kg":                bool(data.get("venta_kg", False)),
            "venta_rollo":             bool(data.get("venta_rollo", False)),
            "venta_metro":             bool(data.get("venta_metro", False)),
            "es_materia_prima":        False,
            "materia_prima_id":        materia_prima_id,
            "derivado_ancho":          ancho if usa_formula else None,
            "derivado_largo":          largo if usa_formula else None,
            "derivado_calibre":        calibre if usa_formula else None,
            "derivado_constante":      constante if usa_formula else None,
            "derivado_sellado":        sellado if usa_formula else None,
            "presentaciones_extra":    presentaciones_extra,
            "activo":                  True,
            "nota_descuento":          data.get("notas") or "",
            "ultima_factura":          ultima_factura_val,
            "ultima_fecha":            ultima_fecha_val,
        }

        res = sb.table("productos").insert(insert_payload).execute()
        nuevo = res.data[0] if res.data else {}

        return {
            "ok": True,
            "producto": nuevo,
            "costo_unidad_sin_iva": costo_unidad,
            "costo_paquete_sin_iva": cp,
            "costo_caja_sin_iva": cc,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    try:
        data     = await file.read()
        filename = file.filename or ""
        ext      = filename.rsplit(".", 1)[-1].lower()
        texto    = ""

        if ext in ("xlsx", "xls"):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True)
                filas = []
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        fila = [str(c) if c is not None else "" for c in row]
                        if any(f.strip() for f in fila):
                            filas.append(" | ".join(fila))
                texto = "\n".join(filas)
            except Exception as e:
                texto = f"Error leyendo Excel: {e}"

        elif ext in ("docx", "doc"):
            try:
                import docx
                doc = docx.Document(io.BytesIO(data))
                texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        texto += "\n" + " | ".join(c.text for c in row.cells)
            except Exception as e:
                texto = f"Error leyendo Word: {e}"

        else:
            texto = data.decode("utf-8", errors="ignore")

        return {"texto": texto[:15000]}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/extract-lista")
async def extract_lista(file: UploadFile = File(...)):
    import base64, google.generativeai as genai

    try:
        GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")
        if not GOOGLE_API_KEY:
            raise HTTPException(status_code=500, detail="GOOGLE_API_KEY no configurada")

        genai.configure(api_key=GOOGLE_API_KEY)
        model = genai.GenerativeModel("gemini-1.5-flash")

        data     = await file.read()
        filename = file.filename or ""
        ext      = filename.rsplit(".", 1)[-1].lower()

        prompt = (
            "Eres un asistente que extrae listas de precios de proveedores. "
            "Analiza este documento y extrae TODOS los productos con sus precios. "
            "Los precios pueden usar punto como separador de miles (ej: 23.243 = 23243 pesos colombianos). "
            "Devuelve los precios como numeros enteros sin puntos ni comas. "
            "Responde SOLO con JSON valido, sin texto adicional, sin backticks: "
            '{"productos":[{"nombre":"nombre del producto","sku":"codigo si existe sino vacio","precio":12500,"unidad":"unidad de medida"}]} '
            "Solo incluye productos reales con precios numericos. Ignora encabezados, totales y notas."
        )

        if ext in ("jpg", "jpeg", "png", "webp"):
            import PIL.Image
            img = PIL.Image.open(io.BytesIO(data))
            response = model.generate_content([prompt, img])
        elif ext == "pdf":
            pdf_part = {"mime_type": "application/pdf", "data": base64.b64encode(data).decode()}
            response = model.generate_content([prompt, pdf_part])
        else:
            texto = ""
            if ext in ("xlsx", "xls"):
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True)
                    filas = []
                    for sheet in wb.worksheets:
                        for row in sheet.iter_rows(values_only=True):
                            fila = [str(c) if c is not None else "" for c in row]
                            if any(f.strip() for f in fila):
                                filas.append(" | ".join(fila))
                    texto = "\n".join(filas)
                except Exception as e:
                    texto = f"Error: {e}"
            elif ext in ("docx", "doc"):
                try:
                    import docxlib
                    doc = docxlib.Document(io.BytesIO(data))
                    lineas = [p.text for p in doc.paragraphs if p.text.strip()]
                    texto = "\n".join(lineas)
                except Exception as e:
                    texto = f"Error: {e}"
            else:
                texto = data.decode("utf-8", errors="ignore")
            response = model.generate_content(prompt + "\n\nContenido:\n" + texto[:10000])

        texto_resp = response.text or ""
        clean      = texto_resp.replace("```json", "").replace("```", "").strip()
        parsed     = json.loads(clean)
        productos  = parsed.get("productos", [])

        return {"productos": productos, "total": len(productos)}

    except json.JSONDecodeError:
        raise HTTPException(status_code=422, detail="No se pudo parsear la respuesta de Gemini")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/config-admin/{clave}")
async def get_config(clave: str):
    try:
        sb  = get_supabase()
        res = sb.table("config_admin").select("valor").eq("clave", clave).maybe_single().execute()
        return {"valor": res.data.get("valor", "") if res.data else ""}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/config-admin")
async def set_config(data: dict):
    try:
        sb = get_supabase()
        sb.table("config_admin").upsert({"clave": data["clave"], "valor": data["valor"]}).execute()
        return {"ok": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/enviar-reporte-mensual")
async def enviar_reporte_mensual(data: dict):
    try:
        sb  = get_supabase()
        mes = data.get("mes", "")

        cfg    = sb.table("config_admin").select("valor").eq("clave", "correo_reporte").maybe_single().execute()
        correo = cfg.data.get("valor", "") if cfg.data else ""
        if not correo:
            raise HTTPException(status_code=400, detail="No hay correo configurado en Admin")

        total_compras       = float(data.get("total_compras", 0))
        total_iva_compras   = float(data.get("total_iva_compras", 0))
        total_rete          = float(data.get("total_rete", 0))
        total_descuentos    = float(data.get("total_descuentos", 0))
        total_arriendos     = float(data.get("total_arriendos", 0))
        ventas              = float(data.get("ventas", 0))
        iva_ventas          = float(data.get("iva_ventas", 0))
        iva_arriendos       = float(data.get("iva_arriendos", 0))
        rete_arriendos      = float(data.get("rete_arriendos", 0))
        rete_total          = float(data.get("rete_total", total_rete + rete_arriendos))
        iva_neto            = float(data.get("iva_neto", 0))
        facturas_pagadas    = int(data.get("facturas_pagadas", 0))
        facturas_pendientes = int(data.get("facturas_pendientes", 0))
        valor_pendiente     = float(data.get("valor_pendiente", 0))

        def fmt(n: float) -> str:
            return f"${int(round(n)):,}".replace(",", ".")

        color_iva = "#f59e0b" if iva_neto > 0 else "#22c55e"

        html = f"""
        <div style="font-family:sans-serif;max-width:600px;margin:0 auto;background:#f4f4f5;padding:32px;border-radius:12px;">
          <div style="background:#C41E2C;padding:20px 24px;border-radius:8px;margin-bottom:24px;">
            <h1 style="color:white;margin:0;font-size:20px;">Resumen Contable {mes}</h1>
            <p style="color:rgba(255,255,255,0.7);margin:6px 0 0;font-size:13px;">Punto Rojo</p>
          </div>

          <div style="background:white;border-radius:8px;padding:20px;margin-bottom:16px;">
            <p style="font-size:11px;color:#888;text-transform:uppercase;letter-spacing:1px;margin:0 0 14px;font-weight:600;">Ventas del mes</p>
            <table width="100%" cellpadding="6" style="font-size:14px;border-collapse:collapse;">
              <tr><td style="color:#444;">Total ventas</td><td align="right" style="font-weight:bold;">{fmt(ventas)}</td></tr>
              <tr><td style="color:#444;">IVA generado (19%)</td><td align="right" style="color:#22c55e;font-weight:bold;">{fmt(iva_ventas)}</td></tr>
            </table>
          </div>

          <div style="background:white;border-radius:8px;padding:20px;margin-bottom:16px;">
            <p style="font-size:11px;color:#888;text-transform:uppercase;letter-spacing:1px;margin:0 0 14px;font-weight:600;">Compras del mes</p>
            <table width="100%" cellpadding="6" style="font-size:14px;border-collapse:collapse;">
              <tr><td style="color:#444;">Total compras (subtotal)</td><td align="right" style="font-weight:bold;">{fmt(total_compras)}</td></tr>
              <tr><td style="color:#444;">IVA descontable</td><td align="right" style="color:#ef4444;">{fmt(total_iva_compras)}</td></tr>
              <tr><td style="color:#444;">Retefuente facturas</td><td align="right" style="color:#ef4444;">{fmt(total_rete)}</td></tr>
              {f'<tr><td style="color:#444;">Retefuente gastos</td><td align="right" style="color:#ef4444;">{fmt(rete_arriendos)}</td></tr>' if rete_arriendos > 0 else ''}
              <tr style="border-top:1px solid #eee;"><td style="color:#111;font-weight:bold;padding-top:8px;">Retefuente total</td><td align="right" style="font-weight:bold;color:#ef4444;padding-top:8px;">{fmt(rete_total)}</td></tr>
              {f'<tr><td style="color:#444;">Descuentos proveedores</td><td align="right" style="color:#22c55e;">- {fmt(total_descuentos)}</td></tr>' if total_descuentos > 0 else ''}
            </table>
          </div>

          <div style="background:white;border-radius:8px;padding:20px;margin-bottom:16px;">
            <p style="font-size:11px;color:#888;text-transform:uppercase;letter-spacing:1px;margin:0 0 14px;font-weight:600;">Arriendos y otros gastos</p>
            <table width="100%" cellpadding="6" style="font-size:14px;border-collapse:collapse;">
              <tr><td style="color:#444;">Total gastos</td><td align="right" style="font-weight:bold;">{fmt(total_arriendos)}</td></tr>
              <tr><td style="color:#444;">IVA gastos</td><td align="right" style="color:#ef4444;">{fmt(iva_arriendos)}</td></tr>
              {f'<tr><td style="color:#444;">Retefuente gastos</td><td align="right" style="color:#ef4444;">{fmt(rete_arriendos)}</td></tr>' if rete_arriendos > 0 else ''}
            </table>
          </div>

          <div style="background:white;border-radius:8px;padding:20px;margin-bottom:16px;border-left:4px solid {color_iva};">
            <p style="font-size:11px;color:#888;text-transform:uppercase;letter-spacing:1px;margin:0 0 14px;font-weight:600;">Resumen IVA</p>
            <table width="100%" cellpadding="6" style="font-size:14px;border-collapse:collapse;">
              <tr><td style="color:#444;">IVA ventas</td><td align="right">{fmt(iva_ventas)}</td></tr>
              <tr><td style="color:#444;">IVA compras</td><td align="right">- {fmt(total_iva_compras)}</td></tr>
              <tr><td style="color:#444;">IVA gastos</td><td align="right">- {fmt(iva_arriendos)}</td></tr>
              <tr style="border-top:1px solid #eee;">
                <td style="color:#111;font-weight:bold;padding-top:10px;">IVA neto a pagar</td>
                <td align="right" style="font-weight:bold;font-size:16px;color:{color_iva};padding-top:10px;">{fmt(iva_neto)}</td>
              </tr>
            </table>
          </div>

          <div style="background:white;border-radius:8px;padding:20px;margin-bottom:24px;">
            <p style="font-size:11px;color:#888;text-transform:uppercase;letter-spacing:1px;margin:0 0 14px;font-weight:600;">Estado de pagos</p>
            <table width="100%" cellpadding="6" style="font-size:14px;border-collapse:collapse;">
              <tr><td style="color:#444;">Facturas pagadas</td><td align="right" style="color:#22c55e;font-weight:bold;">{facturas_pagadas}</td></tr>
              <tr><td style="color:#444;">Facturas pendientes</td><td align="right" style="color:#f59e0b;font-weight:bold;">{facturas_pendientes}</td></tr>
              <tr><td style="color:#444;">Valor pendiente</td><td align="right" style="color:#ef4444;font-weight:bold;">{fmt(valor_pendiente)}</td></tr>
            </table>
          </div>

          <p style="text-align:center;color:#aaa;font-size:11px;margin:0;">Generado desde Costos Punto Rojo</p>
        </div>
        """

        resend.api_key = os.getenv("RESEND_API_KEY", "")
        resend.Emails.send({
            "from":    "Punto Rojo <onboarding@resend.dev>",
            "to":      [correo],
            "subject": f"Resumen Contable {mes} - Punto Rojo",
            "html":    html,
        })

        return {"ok": True, "correo": correo}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── Limpiar tablas ──────────────────────────────────────────────────────────

TABLAS_PERMITIDAS = [
    "historial_descuentos",
    "historial_costos",
    "facturas_revisadas",
    "reportes_revisados",
    "listas_precios_items",
    "listas_precios",
    "productos",
    "facturas_contables",
    "facturas",
]

class LimpiarRequest(BaseModel):
    tablas: list[str]

@app.delete("/limpiar-tablas")
async def limpiar_tablas(req: LimpiarRequest):
    for tabla in req.tablas:
        if tabla not in TABLAS_PERMITIDAS:
            raise HTTPException(status_code=400, detail=f"Tabla no permitida: {tabla}")
    try:
        sb = get_supabase()
        for tabla in req.tablas:
            sb.table(tabla).delete().neq("id", 0).execute()
        return {"ok": True, "tablas_borradas": req.tablas}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── Reporte de costos por email ─────────────────────────────────────────────

@app.post("/enviar-reporte-costos")
async def enviar_reporte_costos(data: dict):
    try:
        sb  = get_supabase()
        fecha = data.get("fecha", "")
        if not fecha:
            raise HTTPException(status_code=400, detail="Se requiere una fecha")

        # Obtener correo configurado
        cfg    = sb.table("config_admin").select("valor").eq("clave", "correo_reporte").maybe_single().execute()
        correo = cfg.data.get("valor", "") if cfg.data else ""
        if not correo:
            raise HTTPException(status_code=400, detail="No hay correo configurado en Admin")

        # Obtener historial del día (solo NUEVO, SUBIO, BAJO)
        desde = f"{fecha}T00:00:00"
        hasta = f"{fecha}T23:59:59"
        hist_res = sb.table("historial_costos").select(
            "producto_id,estado,costo_unidad_anterior,costo_unidad_nuevo,variacion_porcentaje,factura_id,presentacion_facturada,costo_presentacion_facturada,nota"
        ).gte("creada_en", desde).lte("creada_en", hasta).in_("estado", ["NUEVO", "SUBIO", "BAJO"]).order("estado").execute()

        hist = hist_res.data or []

        # Facturas del día (independiente de si hubo cambios de costo)
        facts_dia_res = sb.table("facturas").select(
            "id,numero_factura,nota,revisado,revisado_por"
        ).gte("creada_en", desde).lte("creada_en", hasta).order("id").execute()
        facturas_dia = facts_dia_res.data or []

        if not hist and not facturas_dia:
            raise HTTPException(status_code=400, detail="No hay productos con cambios ni facturas en esa fecha")

        # Obtener productos y facturas relacionados
        prod_ids = list({h["producto_id"] for h in hist})
        fact_ids = list({h["factura_id"] for h in hist if h.get("factura_id")})

        prods_res = sb.table("productos_con_proveedor").select(
            "id,nombre_punto_rojo,sku_interno,categoria,unidades_por_paquete,paquetes_por_caja,unidades_por_caja,"
            "markup_unidad_pct,markup_paquete_pct,markup_caja_pct,markup_millar_pct,markup_kg_pct,markup_rollo_pct,markup_metro_pct,"
            "costo_paquete_sin_iva,costo_caja_sin_iva,venta_unidad,venta_paquete,venta_caja,"
            "venta_millar,venta_kg,venta_rollo,venta_metro,proveedor_nombre"
        ).in_("id", prod_ids).execute()

        facts_res = sb.table("facturas").select("id,numero_factura").in_("id", fact_ids).execute() if fact_ids else type("R", (), {"data": []})()

        prods_map = {p["id"]: p for p in (prods_res.data or [])}
        facts_map = {f["id"]: f for f in (facts_res.data or [])}

        # Combinar datos
        items = []
        for h in hist:
            p = prods_map.get(h["producto_id"], {})
            f = facts_map.get(h.get("factura_id"), {})
            cu = h.get("costo_unidad_nuevo") or 0
            cp = p.get("costo_paquete_sin_iva") or cu * (p.get("unidades_por_paquete") or 1)
            cc = p.get("costo_caja_sin_iva") or cu * (p.get("unidades_por_caja") or 1)
            mu = p.get("markup_unidad_pct") or 40
            mp = p.get("markup_paquete_pct") or 35
            mc = p.get("markup_caja_pct") or 31.58
            mMillar = p.get("markup_millar_pct") or 40
            mKg     = p.get("markup_kg_pct") or 40
            mRollo  = p.get("markup_rollo_pct") or 40
            mMetro  = p.get("markup_metro_pct") or 40
            items.append({
                "estado":      h.get("estado"),
                "nombre":      p.get("nombre_punto_rojo") or "—",
                "sku":         p.get("sku_interno") or "—",
                "categoria":   p.get("categoria") or "—",
                "proveedor":   p.get("proveedor_nombre") or "—",
                "factura":     f.get("numero_factura") or "—",
                "nota":        h.get("nota") or "",
                "costo_ant":   h.get("costo_unidad_anterior") or 0,
                "costo_nuevo": cu,
                "costo_presentacion": h.get("costo_presentacion_facturada") or cu,
                "presentacion": h.get("presentacion_facturada") or "Unidad",
                "variacion":   h.get("variacion_porcentaje") or 0,
                "venta_unidad":  p.get("venta_unidad", True),
                "venta_paquete": p.get("venta_paquete", False),
                "venta_caja":    p.get("venta_caja", False),
                "venta_millar":  p.get("venta_millar", False),
                "venta_kg":      p.get("venta_kg", False),
                "venta_rollo":   p.get("venta_rollo", False),
                "venta_metro":   p.get("venta_metro", False),
                "up": p.get("unidades_por_paquete") or 1,
                "uc": p.get("unidades_por_caja") or 1,
                "precio_unidad":  round(cu * 1.19 * (mu/100) if mu >= 100 else (cu * 1.19) / (1 - mu/100)) if cu else 0,
                "precio_paquete": round(cp * 1.19 * (mp/100) if mp >= 100 else (cp * 1.19) / (1 - mp/100)) if cp else 0,
                "precio_caja":    round(cc * 1.19 * (mc/100) if mc >= 100 else (cc * 1.19) / (1 - mc/100)) if cc else 0,
                "precio_millar":  round(cu * 1000 * 1.19 * (mMillar/100) if mMillar >= 100 else (cu * 1000 * 1.19) / (1 - mMillar/100)) if cu else 0,
                "precio_kg":      round(cu * 1.19 * (mKg/100) if mKg >= 100 else (cu * 1.19) / (1 - mKg/100)) if cu else 0,
                "precio_rollo":   round(cu * 1.19 * (mRollo/100) if mRollo >= 100 else (cu * 1.19) / (1 - mRollo/100)) if cu else 0,
                "precio_metro":   round(cu * 1.19 * (mMetro/100) if mMetro >= 100 else (cu * 1.19) / (1 - mMetro/100)) if cu else 0,
                "mu": mu, "mp": mp, "mc": mc,
            })

        def fmt(n): return f"${int(round(n)):,}".replace(",", ".")
        def fecha_bonita(f): 
            try:
                from datetime import datetime
                return datetime.strptime(f, "%Y-%m-%d").strftime("%d de %B de %Y").replace(
                    "January","enero").replace("February","febrero").replace("March","marzo").replace(
                    "April","abril").replace("May","mayo").replace("June","junio").replace(
                    "July","julio").replace("August","agosto").replace("September","septiembre").replace(
                    "October","octubre").replace("November","noviembre").replace("December","diciembre")
            except: return f

        def estado_color(estado):
            return {"NUEVO": "#C41E2C", "SUBIO": "#ef4444", "BAJO": "#22c55e"}.get(estado, "#888")

        def estado_label(estado):
            return {"NUEVO": "NUEVO", "SUBIO": "↑ SUBIÓ", "BAJO": "↓ BAJÓ"}.get(estado, estado)

        def render_precios(item):
            precios = []
            if item["venta_unidad"]:
                precios.append(f'<td style="padding:4px 8px;"><span style="font-size:10px;color:#888;">Unidad</span><br><strong style="font-size:13px;">{fmt(item["precio_unidad"])}</strong></td>')
            if item["venta_paquete"]:
                precios.append(f'<td style="padding:4px 8px;"><span style="font-size:10px;color:#888;">Paq x{item["up"]}</span><br><strong style="font-size:13px;">{fmt(item["precio_paquete"])}</strong></td>')
            if item["venta_caja"]:
                precios.append(f'<td style="padding:4px 8px;"><span style="font-size:10px;color:#888;">Caja x{item["uc"]}</span><br><strong style="font-size:13px;">{fmt(item["precio_caja"])}</strong></td>')
            if item["venta_millar"]:
                precios.append(f'<td style="padding:4px 8px;"><span style="font-size:10px;color:#888;">Millar</span><br><strong style="font-size:13px;">{fmt(item["precio_millar"])}</strong></td>')
            if item["venta_kg"]:
                precios.append(f'<td style="padding:4px 8px;"><span style="font-size:10px;color:#888;">Kg</span><br><strong style="font-size:13px;">{fmt(item["precio_kg"])}</strong></td>')
            if item["venta_rollo"]:
                precios.append(f'<td style="padding:4px 8px;"><span style="font-size:10px;color:#888;">Rollo</span><br><strong style="font-size:13px;">{fmt(item["precio_rollo"])}</strong></td>')
            if item["venta_metro"]:
                precios.append(f'<td style="padding:4px 8px;"><span style="font-size:10px;color:#888;">Metro</span><br><strong style="font-size:13px;">{fmt(item["precio_metro"])}</strong></td>')
            return "".join(precios)

        # Agrupar por estado
        nuevos   = [i for i in items if i["estado"] == "NUEVO"]
        subieron = [i for i in items if i["estado"] == "SUBIO"]
        bajaron  = [i for i in items if i["estado"] == "BAJO"]

        def render_grupo(titulo, color, grupo):
            if not grupo: return ""
            filas = ""
            for item in grupo:
                variacion_html = ""
                if item["variacion"] != 0 and item["costo_ant"] > 0:
                    arrow = "↑" if item["variacion"] > 0 else "↓"
                    col   = "#ef4444" if item["variacion"] > 0 else "#22c55e"
                    diff_valor = abs(item["costo_nuevo"] - item["costo_ant"])
                    variacion_html = f'<span style="color:{col};font-weight:bold;margin-left:6px;">{arrow}{fmt(diff_valor)} ({abs(item["variacion"])}%)</span>'

                costo_ant_html = ""
                if item["costo_ant"] > 0:
                    costo_ant_html = f'<span style="color:#aaa;font-size:11px;text-decoration:line-through;margin-right:6px;">{fmt(item["costo_ant"])}</span>'

                nota_item_html = ""
                if item.get("nota"):
                    import html as _html_local
                    nota_item_html = f'<p style="font-size:11px;color:#92400e;background:#fffbeb;border-left:3px solid #f59e0b;padding:4px 8px;margin:4px 0 0;border-radius:4px;">⚠️ {_html_local.escape(item["nota"])}</p>'

                filas += f"""
                <tr style="border-bottom:1px solid #f0f0f0;">
                  <td style="padding:10px 12px;vertical-align:top;">
                    <strong style="font-size:13px;color:#111;">{item["nombre"]}</strong><br>
                    <span style="font-size:11px;color:#888;">{item["categoria"]} · {item["proveedor"]}</span><br>
                    <span style="font-size:10px;color:#bbb;">Factura: {item["factura"]}</span>
                    {nota_item_html}
                  </td>
                  <td style="padding:10px 12px;vertical-align:top;white-space:nowrap;">
                    {costo_ant_html}
                    <strong style="font-size:13px;color:#111;">{fmt(item["costo_presentacion"])}</strong>
                    <span style="font-size:11px;color:#888;margin-left:4px;">/ {item["presentacion"]}</span>
                    {variacion_html}
                  </td>
                  <td style="padding:10px 12px;vertical-align:top;">
                    <table cellpadding="0" cellspacing="0"><tr>{render_precios(item)}</tr></table>
                  </td>
                </tr>"""

            return f"""
            <div style="margin-bottom:20px;">
              <div style="background:{color};padding:10px 16px;border-radius:8px 8px 0 0;">
                <strong style="color:white;font-size:13px;">{titulo} ({len(grupo)})</strong>
              </div>
              <table width="100%" cellpadding="0" cellspacing="0" style="background:white;border-radius:0 0 8px 8px;overflow:hidden;border:1px solid #eee;border-top:none;">
                <thead>
                  <tr style="background:#f9f9f9;border-bottom:1px solid #eee;">
                    <th style="padding:8px 12px;text-align:left;font-size:11px;color:#888;font-weight:600;text-transform:uppercase;">Producto</th>
                    <th style="padding:8px 12px;text-align:left;font-size:11px;color:#888;font-weight:600;text-transform:uppercase;">Costo</th>
                    <th style="padding:8px 12px;text-align:left;font-size:11px;color:#888;font-weight:600;text-transform:uppercase;">Precios de venta</th>
                  </tr>
                </thead>
                <tbody>{filas}</tbody>
              </table>
            </div>"""

        resumen_html = f"""
        <div style="display:flex;gap:12px;margin-bottom:24px;flex-wrap:wrap;">
          {"".join(f'<div style="flex:1;min-width:80px;background:white;border-radius:8px;padding:14px;text-align:center;border-top:3px solid {estado_color(e)};"><strong style="font-size:22px;color:{estado_color(e)};">{c}</strong><br><span style="font-size:11px;color:#888;">{l}</span></div>'
            for e, c, l in [("NUEVO", len(nuevos), "Nuevos"), ("SUBIO", len(subieron), "Subieron"), ("BAJO", len(bajaron), "Bajaron")])}
        </div>"""

        # Obtener nota del reporte si existe
        nota_html = ""
        try:
            rev = sb.table("reportes_revisados").select("notas,revisado_por").eq("fecha", fecha).maybe_single().execute()
            if rev.data and rev.data.get("notas"):
                nota_html = f"""
                <div style="background:white;border-radius:8px;padding:16px;margin-bottom:20px;border-left:3px solid #C41E2C;">
                  <p style="font-size:11px;color:#888;font-weight:600;text-transform:uppercase;margin:0 0 6px;">Notas del reporte</p>
                  <p style="font-size:13px;color:#333;margin:0;">{rev.data["notas"]}</p>
                </div>"""
        except Exception:
            pass

        # Facturas del día con estado de revisión
        import html as _html
        if facturas_dia:
            filas_facturas = ""
            for f in facturas_dia:
                if f.get("revisado"):
                    badge = f'<span style="display:inline-block;background:#dcfce7;color:#16a34a;font-size:10px;font-weight:600;padding:3px 8px;border-radius:12px;">✓ Revisado{(" por " + _html.escape(f["revisado_por"])) if f.get("revisado_por") else ""}</span>'
                else:
                    badge = '<span style="display:inline-block;background:#f3f4f6;color:#9ca3af;font-size:10px;font-weight:600;padding:3px 8px;border-radius:12px;">Sin revisar</span>'

                nota_html_factura = ""
                if f.get("nota") and f["nota"].strip():
                    nota_html_factura = f'<p style="font-size:12px;color:#92400e;background:#fffbeb;border-left:3px solid #f59e0b;padding:6px 10px;margin:6px 0 0;border-radius:4px;">⚠️ {_html.escape(f["nota"])}</p>'

                filas_facturas += f"""
                <tr style="border-bottom:1px solid #f0f0f0;">
                  <td style="padding:8px 12px;vertical-align:top;">
                    <strong style="font-size:12px;color:#111;">{_html.escape(f.get("numero_factura") or f"Factura #{f['id']}")}</strong>
                    {nota_html_factura}
                  </td>
                  <td style="padding:8px 12px;vertical-align:top;text-align:right;white-space:nowrap;">{badge}</td>
                </tr>"""

            n_revisadas = sum(1 for f in facturas_dia if f.get("revisado"))
            facturas_dia_html = f"""
            <div style="margin-bottom:20px;">
              <div style="background:#444;padding:10px 16px;border-radius:8px 8px 0 0;">
                <strong style="color:white;font-size:13px;">📄 Facturas del día ({len(facturas_dia)}) · {n_revisadas} revisadas</strong>
              </div>
              <table width="100%" cellpadding="0" cellspacing="0" style="background:white;border-radius:0 0 8px 8px;overflow:hidden;border:1px solid #eee;border-top:none;">
                <tbody>{filas_facturas}</tbody>
              </table>
            </div>"""
        else:
            facturas_dia_html = ""

        html = f"""
        <div style="font-family:sans-serif;max-width:680px;margin:0 auto;background:#f4f4f5;padding:32px;border-radius:12px;">
          <div style="background:#C41E2C;padding:20px 24px;border-radius:8px;margin-bottom:24px;">
            <h1 style="color:white;margin:0;font-size:20px;">Reporte de Costos</h1>
            <p style="color:rgba(255,255,255,0.75);margin:6px 0 0;font-size:13px;">{fecha_bonita(fecha)} · Punto Rojo</p>
          </div>
          {resumen_html}
          {nota_html}
          {facturas_dia_html}
          {render_grupo("🆕 Productos nuevos",   "#C41E2C", nuevos)}
          {render_grupo("📈 Subieron de precio", "#ef4444", subieron)}
          {render_grupo("📉 Bajaron de precio",  "#22c55e", bajaron)}
          <p style="text-align:center;color:#aaa;font-size:11px;margin:0;">Generado desde Costos Punto Rojo</p>
        </div>"""

        resend.api_key = os.getenv("RESEND_API_KEY", "")
        resend.Emails.send({
            "from":    "Punto Rojo <onboarding@resend.dev>",
            "to":      [correo],
            "subject": f"Reporte de Costos {fecha_bonita(fecha)} - Punto Rojo",
            "html":    html,
        })

        return {"ok": True, "correo": correo, "total": len(items)}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── Marcar factura revisada ────────────────────────────────────────────────

@app.post("/marcar-revisado-factura")
async def marcar_revisado_factura(data: dict):
    try:
        sb = get_supabase()
        factura_id = data.get("factura_id")
        revisado   = bool(data.get("revisado", True))
        revisado_por = data.get("revisado_por") or "Usuario"

        if not factura_id:
            raise HTTPException(status_code=400, detail="Se requiere factura_id")

        update_data = {"revisado": revisado}
        if revisado:
            update_data["revisado_por"] = revisado_por
            update_data["revisado_en"]  = datetime.utcnow().isoformat()
        else:
            update_data["revisado_por"] = None
            update_data["revisado_en"]  = None

        sb.table("facturas").update(update_data).eq("id", factura_id).execute()
        return {"ok": True}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── Unificar productos duplicados ───────────────────────────────────────────

class UnificarProductosRequest(BaseModel):
    sobrevive_id: int
    duplicado_id: int

COLS_COSTO_VIGENTE = (
    "id,sku_interno,nombre_punto_rojo,costo_unidad_sin_iva,costo_paquete_sin_iva,costo_caja_sin_iva,"
    "ultima_factura,ultima_fecha,nota_descuento,descuento_pct_factura,descuento_aplicado,precio_factura_base"
)

@app.post("/unificar-productos")
async def unificar_productos(data: UnificarProductosRequest, request: Request):
    """Fusiona dos productos duplicados: mueve el historial de costos y
    descuentos del duplicado hacia el sobreviviente, y borra el duplicado.
    El costo vigente del sobreviviente se actualiza al de la factura MAS
    RECIENTE entre los dos -- sin importar cual de los dos se haya elegido
    como sobreviviente -- para no perder un costo mas nuevo que haya llegado
    al que se termina borrando. El markup NO se toca (es una decision de
    precio de venta, no algo atado a la fecha de la factura)."""
    try:
        if data.sobrevive_id == data.duplicado_id:
            raise HTTPException(status_code=400, detail="Los dos productos deben ser distintos")

        sb = get_supabase()
        sobrevive = sb.table("productos").select(COLS_COSTO_VIGENTE).eq("id", data.sobrevive_id).maybe_single().execute()
        duplicado = sb.table("productos").select(COLS_COSTO_VIGENTE).eq("id", data.duplicado_id).maybe_single().execute()
        if not sobrevive.data or not duplicado.data:
            raise HTTPException(status_code=404, detail="Producto no encontrado")

        costo_actualizado = False
        fecha_sobrevive = sobrevive.data.get("ultima_fecha") or ""
        fecha_duplicado = duplicado.data.get("ultima_fecha") or ""
        if fecha_duplicado and fecha_duplicado > fecha_sobrevive:
            sb.table("productos").update({
                "costo_unidad_sin_iva":   duplicado.data["costo_unidad_sin_iva"],
                "costo_paquete_sin_iva":  duplicado.data["costo_paquete_sin_iva"],
                "costo_caja_sin_iva":     duplicado.data["costo_caja_sin_iva"],
                "ultima_factura":         duplicado.data["ultima_factura"],
                "ultima_fecha":           duplicado.data["ultima_fecha"],
                "nota_descuento":         duplicado.data["nota_descuento"],
                "descuento_pct_factura":  duplicado.data["descuento_pct_factura"],
                "descuento_aplicado":     duplicado.data["descuento_aplicado"],
                "precio_factura_base":    duplicado.data["precio_factura_base"],
            }).eq("id", data.sobrevive_id).execute()
            costo_actualizado = True

        sb.table("historial_costos").update({"producto_id": data.sobrevive_id}).eq("producto_id", data.duplicado_id).execute()
        sb.table("historial_descuentos").update({"producto_id": data.sobrevive_id}).eq("producto_id", data.duplicado_id).execute()
        sb.table("productos").delete().eq("id", data.duplicado_id).execute()

        usuario_email  = request.headers.get("X-Usuario-Email", "")
        usuario_nombre = request.headers.get("X-Usuario-Nombre", "")
        registrar_auditoria(sb,
            accion="UNIFICAR_PRODUCTOS", entidad="productos", entidad_id=str(data.sobrevive_id),
            descripcion=(
                f"Producto {duplicado.data['sku_interno']} ({duplicado.data['nombre_punto_rojo']}) "
                f"fusionado dentro de {sobrevive.data['sku_interno']} ({sobrevive.data['nombre_punto_rojo']})"
                + (" -- costo actualizado al mas reciente" if costo_actualizado else "")
            ),
            usuario_email=usuario_email, usuario_nombre=usuario_nombre,
            datos_anteriores={"sobrevive": sobrevive.data, "duplicado": duplicado.data},
        )

        return {"ok": True, "sobrevive_id": data.sobrevive_id, "duplicado_eliminado": data.duplicado_id, "costo_actualizado": costo_actualizado}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── Eliminar factura ────────────────────────────────────────────────────────

@app.delete("/eliminar-factura/{factura_id}")
async def eliminar_factura(factura_id: int, request: Request):
    try:
        sb = get_supabase()

        # 1. Obtener datos de la factura
        fac = sb.table("facturas").select("id,numero_factura").eq("id", factura_id).maybe_single().execute()
        if not fac.data:
            raise HTTPException(status_code=404, detail="Factura no encontrada")
        numero_factura = fac.data["numero_factura"]

        # 2. Obtener productos afectados por esta factura
        hist = sb.table("historial_costos").select("producto_id").eq("factura_id", factura_id).execute()
        prod_ids = list({h["producto_id"] for h in (hist.data or [])})

        # 3. Borrar historial_descuentos de esa factura
        sb.table("historial_descuentos").delete().eq("factura_id", factura_id).execute()

        # 4. Borrar historial_costos de esa factura
        sb.table("historial_costos").delete().eq("factura_id", factura_id).execute()

        # 5. Borrar notas_credito relacionadas (y capturar proveedor/fecha para
        # recalcular la retefuente del grupo después de borrar, por si esta factura
        # formaba parte de una suma con otras del mismo proveedor/día)
        fc = sb.table("facturas_contables").select("id,proveedor,fecha_factura").eq("numero_factura", numero_factura).execute()
        proveedor_fc = None
        fecha_fc = None
        if fc.data:
            sb.table("notas_credito").delete().eq("factura_contable_id", fc.data[0]["id"]).execute()
            proveedor_fc = fc.data[0].get("proveedor")
            fecha_fc = fc.data[0].get("fecha_factura")

        # 6. Borrar factura contable
        sb.table("facturas_contables").delete().eq("numero_factura", numero_factura).execute()

        # 6b. Recalcular retefuente de las facturas restantes del mismo proveedor/fecha
        # (si esta factura sumaba al total que activaba o desactivaba la base mínima)
        if proveedor_fc and fecha_fc:
            recalcular_retefuente_grupo(sb, proveedor_fc, fecha_fc, "SI")

        # 7. Restaurar costo anterior en productos (desde historial previo).
        # Si un producto se queda SIN ningún historial restante, significa que esta
        # factura era su única referencia real (ej. factura de prueba) — en vez de
        # dejarlo con el costo contaminado para siempre, se desactiva (igual que se
        # hace con productos corruptos por SKU duplicado). La próxima factura real
        # de ese producto creará uno nuevo y limpio.
        productos_desactivados = []
        for prod_id in prod_ids:
            try:
                ultimo = sb.table("historial_costos").select(
                    "costo_unidad_nuevo,costo_presentacion_facturada"
                ).eq("producto_id", prod_id).order("id", desc=True).limit(1).execute()
                if ultimo.data:
                    sb.table("productos").update({
                        "costo_unidad_sin_iva": ultimo.data[0]["costo_unidad_nuevo"],
                    }).eq("id", prod_id).execute()
                else:
                    sb.table("productos").update({"activo": False}).eq("id", prod_id).execute()
                    productos_desactivados.append(prod_id)
            except Exception:
                pass

        # 8. Borrar la factura principal
        sb.table("facturas").delete().eq("id", factura_id).execute()

        # Registrar auditoría
        usuario_email  = request.headers.get("X-Usuario-Email", "")
        usuario_nombre = request.headers.get("X-Usuario-Nombre", "")
        registrar_auditoria(sb,
            accion="ELIMINAR_FACTURA", entidad="facturas", entidad_id=str(factura_id),
            descripcion=f"Factura {numero_factura} eliminada — {len(prod_ids)} productos afectados",
            usuario_email=usuario_email, usuario_nombre=usuario_nombre,
        )

        return {
            "ok": True,
            "numero_factura": numero_factura,
            "productos_afectados": len(prod_ids),
            "productos_desactivados": productos_desactivados,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── Sincronizar forma de pago proveedor → facturas ──────────────────────────

@app.post("/sincronizar-forma-pago")
async def sincronizar_forma_pago(data: dict):
    try:
        sb = get_supabase()
        proveedor_nombre = data.get("proveedor_nombre", "")
        forma_pago        = data.get("forma_pago", "")
        descuento_pct     = data.get("descuento_pct")
        aplica_retefuente = data.get("aplica_retefuente")

        if not proveedor_nombre:
            raise HTTPException(status_code=400, detail="Se requiere proveedor_nombre")

        # Obtener lista exacta de nombres en facturas_contables para este proveedor
        facturas = sb.table("facturas_contables").select("proveedor").execute()
        nombres_unicos = list({f["proveedor"] for f in (facturas.data or []) if f.get("proveedor")})

        # Encontrar los que hacen match por palabras clave
        nombre_lower = proveedor_nombre.lower()
        palabras_clave = [p for p in nombre_lower.split() if len(p) > 3]

        nombres_match = []
        for nombre in nombres_unicos:
            nombre_f_lower = nombre.lower()
            if any(p in nombre_f_lower for p in palabras_clave):
                nombres_match.append(nombre)

        # Obtener parámetros de retefuente activos
        params_rete = sb.table("parametros_retefuente").select(
            "porcentaje,base_minima"
        ).eq("aplica_a", "COMPRAS").eq("activo", True).limit(1).execute()
        pct_rete = float(params_rete.data[0].get("porcentaje") or 2.5) if params_rete.data else 2.5
        base_min = float(params_rete.data[0].get("base_minima") or 1148000) if params_rete.data else 1148000

        # Armar el update con los campos que vienen
        for nombre in nombres_match:
            update_data = {}
            if forma_pago:
                update_data["forma_pago"] = forma_pago

            # Obtener todas las facturas del proveedor
            facts = sb.table("facturas_contables").select(
                "id,subtotal,iva,retefuente,valor_descuento,fecha_factura"
            ).eq("proveedor", nombre).execute()

            if descuento_pct is not None:
                desc = float(descuento_pct)
                update_data["descuento_pct"] = desc
                for f in (facts.data or []):
                    subtotal    = float(f.get("subtotal") or 0)
                    iva         = float(f.get("iva") or 0)
                    retefuente  = float(f.get("retefuente") or 0)
                    valor_desc  = round(subtotal * desc / 100, 2)
                    valor_pagar = max(0.0, subtotal + iva - valor_desc - retefuente)
                    sb.table("facturas_contables").update({
                        "descuento_pct":   desc,
                        "valor_descuento": valor_desc,
                        "valor_a_pagar":   valor_pagar,
                    }).eq("id", f["id"]).execute()

            if aplica_retefuente is not None:
                # Agrupar facturas por fecha y recalcular retefuente
                fechas = list({f["fecha_factura"] for f in (facts.data or []) if f.get("fecha_factura")})
                for fecha in fechas:
                    facts_fecha = [f for f in (facts.data or []) if f.get("fecha_factura") == fecha]
                    total_subtotal = sum(float(f.get("subtotal") or 0) for f in facts_fecha)
                    aplica = aplica_retefuente == "SI" and total_subtotal >= base_min
                    for f in facts_fecha:
                        subtotal   = float(f.get("subtotal") or 0)
                        iva        = float(f.get("iva") or 0)
                        valor_desc = float(f.get("valor_descuento") or 0)
                        rete_f     = round(subtotal * pct_rete / 100, 2) if aplica else 0.0
                        valor_pagar = max(0.0, subtotal + iva - valor_desc - rete_f)
                        sb.table("facturas_contables").update({
                            "aplica_retefuente": aplica_retefuente,
                            "retefuente":        rete_f,
                            "valor_a_pagar":     valor_pagar,
                        }).eq("id", f["id"]).execute()

            if update_data and aplica_retefuente is None and descuento_pct is None:
                sb.table("facturas_contables").update(update_data).eq("proveedor", nombre).execute()

        return {"ok": True, "proveedor": proveedor_nombre, "nombres_match": nombres_match}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/registrar-auditoria")
async def registrar_auditoria_endpoint(data: dict):
    """Endpoint para que el frontend registre acciones de auditoría."""
    try:
        sb = get_supabase()
        registrar_auditoria(sb,
            accion=data.get("accion", ""),
            entidad=data.get("entidad", ""),
            entidad_id=data.get("entidad_id", ""),
            descripcion=data.get("descripcion", ""),
            usuario_email=data.get("usuario_email", ""),
            usuario_nombre=data.get("usuario_nombre", ""),
            datos_anteriores=data.get("datos_anteriores"),
            datos_nuevos=data.get("datos_nuevos"),
        )
        return {"ok": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
